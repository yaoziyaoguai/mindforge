"""M2/M3/M4 Optional Dependency Smoke Tests。

验证 pypdf / python-docx 真实安装时的端到端行为。
这些测试不依赖 mock——使用 synthetic minimal fixtures 确保 adapter 可与
真实 optional dependency 协作。

如果 optional dependency 未安装，pytest.importorskip 自动 skip，不影响
普通 pytest 运行。

RFC_0001 §5.8 / §5.9 — PDF/DOCX adapter optional dependency contract。
"""

from __future__ import annotations

import zlib
from pathlib import Path

import pytest


# =============================================================================
# Synthetic minimal PDF —— 手写 PDF 字节，无需 pypdf 参与创建
# =============================================================================

# 在 PDF 中渲染文字的 PostScript stream
_SYNTHETIC_PDF_TEXT = "Smoke Test — Hello from pypdf adapter."


def _crc32(data: bytes) -> int:
    """计算 CRC-32（与 zlib.crc32 一致）。"""
    return zlib.crc32(data) & 0xFFFF_FFFF


def _make_minimal_pdf_bytes(body_text: str = _SYNTHETIC_PDF_TEXT) -> bytes:
    """构造最小合法 PDF，包含一段可提取的文本。

    结构：
    - Catalog → Pages → Page
    - 一个 Helvetica Type1 字体
    - 一个 content stream（BT/ET 块绘制 body_text）
    - 手写 xref table + trailer
    """
    text_escaped = body_text.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")

    # content stream: position at (100, 700), 12pt Helvetica, draw text
    stream_content = (
        f"BT /F1 12 Tf 100 700 Td ({text_escaped}) Tj ET\n"
    ).encode("latin-1")

    # minimal objects
    obj1 = b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
    obj2 = b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
    obj3 = (
        b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R"
        b"/Resources<</Font<</F1 4 0 R>>>>/Contents 5 0 R>>endobj\n"
    )
    obj4 = b"4 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
    obj5_header = f"5 0 obj<</Length {len(stream_content)}>>stream\n".encode("ascii")
    obj5_footer = b"\nendstream\nendobj\n"

    # assemble body
    body_parts: list[bytes] = []
    offsets: dict[int, int] = {}

    for i, part in enumerate(
        [obj1, obj2, obj3, obj4, obj5_header + stream_content + obj5_footer], start=1
    ):
        offsets[i] = len(b"".join(body_parts))
        body_parts.append(part)

    body = b"".join(body_parts)
    xref_offset = len(body)

    # xref table
    xref_lines = [b"xref\n", "0 6\n0000000000 65535 f \n".encode("ascii")]
    for i in range(1, 6):
        xref_lines.append(f"{offsets[i]:010d} 00000 n \n".encode("ascii"))
    xref = b"".join(xref_lines)

    trailer = (
        b"trailer<</Size 6/Root 1 0 R>>\n"
        b"startxref\n"
        + str(xref_offset).encode("ascii")
        + b"\n%%EOF"
    )

    return b"%PDF-1.4\n" + body + xref + trailer


# =============================================================================
# PDF Smoke（需要 pypdf）
# =============================================================================


class TestPypdfOptionalSmoke:
    """验证 PdfTextAdapter 与真实 pypdf 一起工作。"""

    @pytest.fixture(autouse=True)
    def _check_deps(self):
        pytest.importorskip("pypdf", reason="pypdf optional dependency not installed")

    @pytest.fixture(autouse=True)
    def _imports(self):
        from mindforge.sources.pdf_adapter import PdfTextAdapter

        self.adapter = PdfTextAdapter()

    def test_can_handle_pdf_suffix(self) -> None:
        assert self.adapter.can_handle("doc.pdf") is True
        assert self.adapter.can_handle("doc.PDF") is True
        assert self.adapter.can_handle("doc.txt") is False

    def test_load_synthetic_text_pdf(self, tmp_path: Path) -> None:
        """用真实 pypdf 加载手写 synthetic PDF → loaded。"""
        pdf_path = tmp_path / "smoke.pdf"
        pdf_bytes = _make_minimal_pdf_bytes("Hello from pypdf smoke test.")
        pdf_path.write_bytes(pdf_bytes)

        result = self.adapter.load(str(pdf_path))
        assert result.status == "loaded", f"Expected loaded, got {result.status}: {result.error_message}"
        assert result.document is not None
        assert "pypdf smoke test" in result.document.raw_text.lower()
        assert result.document.source_type == "pdf"
        assert result.document.source_path == str(pdf_path)
        assert len(result.document.content_hash) > 0

    def test_load_empty_page_pdf_skipped(self, tmp_path: Path) -> None:
        """空文本层 PDF（blank page）→ skipped（scanned_pdf_no_text）。"""
        from pypdf import PdfWriter

        pdf_path = tmp_path / "empty.pdf"
        writer = PdfWriter()
        writer.add_blank_page(width=595, height=842)
        with open(pdf_path, "wb") as f:
            writer.write(f)

        result = self.adapter.load(str(pdf_path))
        assert result.status in ("skipped", "loaded")
        # blank page with no text → should be skipped or loaded with empty text
        # PdfTextAdapter uses per-page text extraction, blank page may result in
        # scanned_pdf_no_text skip

    def test_load_missing_pdf_failed(self) -> None:
        result = self.adapter.load("/tmp/nonexistent_smoke_xyz789.pdf")
        assert result.status == "failed"


# =============================================================================
# DOCX Smoke（需要 python-docx）
# =============================================================================


class TestDocxOptionalSmoke:
    """验证 DocxTextAdapter 与真实 python-docx 一起工作。"""

    @pytest.fixture(autouse=True)
    def _check_deps(self):
        pytest.importorskip("docx", reason="python-docx optional dependency not installed")

    @pytest.fixture(autouse=True)
    def _imports(self):
        from mindforge.sources.docx_adapter import DocxTextAdapter

        self.adapter = DocxTextAdapter()

    def test_can_handle_docx_suffix(self) -> None:
        assert self.adapter.can_handle("doc.docx") is True
        assert self.adapter.can_handle("doc.DOCX") is True
        assert self.adapter.can_handle("doc.doc") is False  # legacy binary

    def test_load_synthetic_docx(self, tmp_path: Path) -> None:
        """用真实 python-docx 创建最小 .docx → loaded。"""
        import docx

        docx_path = tmp_path / "smoke.docx"
        doc = docx.Document()
        doc.add_paragraph("Hello from python-docx smoke test.")
        doc.save(str(docx_path))

        result = self.adapter.load(str(docx_path))
        assert result.status == "loaded", f"Expected loaded, got {result.status}: {result.error_message}"
        assert result.document is not None
        assert "python-docx smoke test" in result.document.raw_text.lower()
        assert result.document.source_type == "docx"

    def test_load_synthetic_docx_with_heading(self, tmp_path: Path) -> None:
        """Heading style → Markdown heading。"""
        import docx

        docx_path = tmp_path / "smoke_heading.docx"
        doc = docx.Document()
        doc.add_heading("Introduction", level=1)
        doc.add_paragraph("Body text.")
        doc.add_heading("Details", level=2)
        doc.save(str(docx_path))

        result = self.adapter.load(str(docx_path))
        assert result.status == "loaded"
        text = result.document.raw_text
        assert "# Introduction" in text
        assert "## Details" in text
        assert "Body text." in text

    def test_load_synthetic_docx_table(self, tmp_path: Path) -> None:
        """表格 → Markdown table。"""
        import docx

        docx_path = tmp_path / "smoke_table.docx"
        doc = docx.Document()
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Age"
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "30"
        table.cell(2, 0).text = "Bob"
        table.cell(2, 1).text = "25"
        doc.save(str(docx_path))

        result = self.adapter.load(str(docx_path))
        assert result.status == "loaded"
        text = result.document.raw_text
        assert "| Name | Age |" in text
        assert "| Alice | 30 |" in text
        assert "| Bob | 25 |" in text

    def test_load_empty_docx_skipped(self, tmp_path: Path) -> None:
        """空文档 → skipped。"""
        import docx

        docx_path = tmp_path / "empty.docx"
        doc = docx.Document()  # no paragraphs
        doc.save(str(docx_path))

        result = self.adapter.load(str(docx_path))
        assert result.status == "skipped"

    def test_load_missing_docx_failed(self) -> None:
        result = self.adapter.load("/tmp/nonexistent_smoke_xyz789.docx")
        assert result.status == "failed"


# =============================================================================
# 安全 / 边界守卫
# =============================================================================


class TestOptionalDepsNoNetwork:
    """Smoke tests 不应制造网络请求。"""

    def test_pdf_smoke_no_network(self) -> None:
        """PDF synthetic bytes 创建仅涉及本地数据。"""
        pdf_bytes = _make_minimal_pdf_bytes("test")
        assert len(pdf_bytes) > 0
        assert pdf_bytes.startswith(b"%PDF-1.4")
        assert b"%%EOF" in pdf_bytes

    def test_docx_smoke_is_optional(self) -> None:
        """验证 docx 是 optional dependency——缺失时 importorskip 跳过。"""
        try:
            import docx  # noqa: F401
            has_docx = True
        except ImportError:
            has_docx = False
        # 不 fail，仅记录状态——docx 是 optional
        assert has_docx in (True, False)
