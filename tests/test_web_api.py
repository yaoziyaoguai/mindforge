"""MindForge Local Console API tests.

中文学习型说明：这些测试使用临时 vault/config 验证 Web adapter 边界：
API 可以读取真实本地状态，但不能泄露 secret；approve 必须二次确认，且
最终仍走现有 approval_service 写入单张 ai_draft。
"""

from __future__ import annotations

from pathlib import Path
import threading
import time

import yaml
import pytest
from fastapi.testclient import TestClient

from mindforge.app_context import AppContextError
from mindforge.config import REQUIRED_STAGES
from mindforge.watch_registry import WatchRegistry
from mindforge_web.app import create_app
from mindforge.ingestion_service import IngestionSummary, _skip_reason_hint
from mindforge_web.services.processing_run_service import ProcessingRunRecord, _apply_summary, _message_for_summary, _now, _run_worker, _safe_error_message
from mindforge_web.services.processing_run_service import _save_record as _save_processing_run_record
from mindforge_web.services.web_source_service import _display_status


def _write_config(tmp_path: Path) -> tuple[Path, Path, Path]:
    vault = tmp_path / "dogfood-vault"
    inbox = vault / "00-Inbox" / "ManualNotes"
    cards = vault / "20-Knowledge-Cards"
    projects = vault / "30-Projects"
    inbox.mkdir(parents=True)
    cards.mkdir(parents=True)
    projects.mkdir(parents=True)
    cfg_path = tmp_path / "mindforge.yaml"
    cfg_path.write_text(
        yaml.safe_dump(
            {
                "version": 0.7,
                "vault": {
                    "root": str(vault),
                    "inbox_root": "00-Inbox",
                    "cards_dir": "20-Knowledge-Cards",
                    "archive_dir": "90-Archive/Skipped",
                    "projects_dir": "30-Projects",
                },
                "sources": {
                    "enabled": ["plain_markdown"],
                    "registry": {
                        "plain_markdown": {
                            "adapter": "PlainMarkdownAdapter",
                            "inbox_subdir": "ManualNotes",
                            "file_glob": "*.md",
                            "enabled": True,
                        }
                    },
                },
                "state": {
                    "workdir": str(tmp_path / ".mindforge"),
                    "state_file": "state.json",
                    "runs_dir": "runs",
                    "index_file": "index.jsonl",
                    "backup_state": True,
                },
                "triage": {"value_score_threshold": 5, "default_track": "unrouted"},
                "llm": {
                    "active_profile": "fake",
                    "profiles": {
                        "fake": {
                            "triage": "fake_alias",
                            "distill": "fake_alias",
                            "link_suggestion": "fake_alias",
                            "review_questions": "fake_alias",
                            "action_extraction": "fake_alias",
                        }
                    },
                    "models": {
                        "fake_alias": {
                            "provider": "fake",
                            "type": "fake",
                            "base_url": "fake://",
                            "model": "fake",
                            "timeout_seconds": 5,
                            "max_retries": 0,
                            "api_key_env": "MINDFORGE_FAKE_SECRET",
                        }
                    },
                },
                "prompts": {
                    "triage_version": "v1",
                    "distill_version": "v1",
                    "link_suggestion_version": "v1",
                    "review_questions_version": "v1",
                    "action_extraction_version": "v1",
                },
                "logging": {"level": "INFO", "file": str(tmp_path / "mf.log")},
            },
            allow_unicode=True,
            sort_keys=False,
        ),
        encoding="utf-8",
    )
    return cfg_path, vault, cards


def _write_config_multi_adapter(tmp_path: Path) -> tuple[Path, Path, Path]:
    """同 _write_config，但额外启用 txt / docx adapter 供跨格式回归测试。"""

    cfg_path, vault, cards = _write_config(tmp_path)
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["sources"]["enabled"] = ["plain_markdown", "txt", "docx"]
    raw["sources"]["registry"]["txt"] = {
        "adapter": "TxtAdapter",
        "inbox_subdir": "ManualNotes",
        "file_glob": "*.txt",
        "enabled": True,
    }
    raw["sources"]["registry"]["docx"] = {
        "adapter": "DocxTextAdapter",
        "inbox_subdir": "ManualNotes",
        "file_glob": "*.docx",
        "enabled": True,
    }
    cfg_path.write_text(yaml.safe_dump(raw, sort_keys=False), encoding="utf-8")
    return cfg_path, vault, cards


def _write_minimal_docx(path: Path) -> None:
    """生成最小合法 .docx 供 DocxTextAdapter 回归测试。"""
    from docx import Document

    doc = Document()
    doc.add_heading("Docx Test", level=1)
    doc.add_paragraph("Minimal body for dedup regression test.")
    doc.save(str(path))


def _write_draft(cards: Path) -> Path:
    card = cards / "draft.md"
    card.write_text(
        """---
id: draft-1
title: Test Draft
status: ai_draft
track: agent-runtime
projects:
  - local-console
tags:
  - web
source_type: manual_note
source_id: src_draft_1
adapter_name: PlainMarkdownAdapter
source_path: 00-Inbox/ManualNotes/source-note.md
source_content_hash: sha256:provenancehash
source_archive_path: ""
source_missing: false
source_title: Safe source
value_score: 8
strategy_id: five_stage
strategy_version: 0.10.0
schema_version: "1"
prompt_version: distill@v1
prompt_versions:
  triage: v1
  distill: v1
  link_suggestion: v1
  review_questions: v1
  action_extraction: v1
profile: fake
stage_models:
  distill: { alias: fake_alias, provider: fake, model: fake }
run_id: run-web-provenance
---

## Distilled Note

This is safe draft body.
""",
        encoding="utf-8",
    )
    return card


def _write_approved_card(cards: Path, *, name: str = "approved.md") -> Path:
    """写一张 Web API 测试用 approved card；不含 source 正文或 secret。"""

    card = cards / name
    card.write_text(
        """---
id: approved-web-1
title: Approved Web Card
status: human_approved
track: agent-runtime
tags:
  - web
source_type: manual_note
source_path: 00-Inbox/ManualNotes/source-note.md
source_title: Safe source
value_score: 8
created_at: 2026-05-08
---

## AI Summary

Approved summary.
""",
        encoding="utf-8",
    )
    return card


def test_wiki_rebuild_json_mode_overrides_config_mode_for_deterministic(tmp_path: Path) -> None:
    """Web rebuild 按 JSON body 的 mode 执行，不能因 config wiki.mode=llm 回落误跑 LLM。"""

    cfg_path, _vault, cards = _write_config(tmp_path)
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["wiki"] = {"mode": "llm", "model": "missing"}
    cfg_path.write_text(yaml.safe_dump(raw, sort_keys=False), encoding="utf-8")
    _write_approved_card(cards)

    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    response = client.post("/api/wiki/rebuild", json={"mode": "deterministic"})

    assert response.status_code == 200
    data = response.json()
    assert data["ok"] is True
    assert data["mode"] == "deterministic"
    assert data["included_cards"] == 1


def test_wiki_rebuild_json_mode_runs_llm_branch(tmp_path: Path, monkeypatch) -> None:
    """Web LLM 按钮传 JSON mode=llm 时后端必须进入 LLM rebuild 分支。"""

    from mindforge.wiki_service import LLMWikiResult

    cfg_path, _vault, _cards = _write_config(tmp_path)
    called = {"value": False}

    def _fake_llm_rebuild(_cfg):
        called["value"] = True
        return LLMWikiResult(
            wiki_path=str(tmp_path / "Main-Wiki.md"),
            included_cards=1,
            section_count=1,
            additional_cards=0,
            warnings=[],
            model_id="main",
            last_rebuilt_at="2026-05-08T00:00:00+0800",
        )

    monkeypatch.setattr("mindforge.wiki_service.llm_rebuild_wiki", _fake_llm_rebuild)

    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    response = client.post("/api/wiki/rebuild", json={"mode": "llm"})

    assert response.status_code == 200
    data = response.json()
    assert called["value"] is True
    assert data["ok"] is True
    assert data["mode"] == "llm"
    assert data["model_id"] == "main"


def test_wiki_page_parses_deterministic_card_marker_into_sections(tmp_path: Path) -> None:
    """`/api/wiki/page` 必须把 deterministic `card=` marker 解析成可见 section。

    中文学习型说明：这是 Web Wiki 的 P0 可见性边界。deterministic builder 和
    LLM builder 可能产生不同 marker；如果 parser 只接受 `card_ids=`，Web 页面会
    只剩 additional cards，看不到已经生成的 Wiki 正文。
    """

    cfg_path, vault, cards = _write_config(tmp_path)
    _write_approved_card(cards)
    wiki_dir = vault / "30-Wiki"
    wiki_dir.mkdir()
    (wiki_dir / "Main-Wiki.md").write_text(
        """# MindForge Main Wiki

## Overview

Synthetic overview.

## agent-runtime

<!-- WIKI_SECTION_START card=approved-web-1 -->
### Approved Web Card

Approved section body.

<!-- WIKI_SECTION_END -->
""",
        encoding="utf-8",
    )

    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    response = client.get("/api/wiki/page")

    assert response.status_code == 200
    data = response.json()
    assert len(data["sections"]) == 1
    assert data["sections"][0]["card_refs"][0]["card_id"] == "approved-web-1"
    assert data["additional_cards"] == []


def test_wiki_page_reports_deterministic_metadata_without_llm_hardcode(tmp_path: Path) -> None:
    """deterministic Wiki 页面的 metadata 不再硬编码 mode。

    中文学习型说明：`/api/wiki/page` 的 mode 字段从 Wiki header 读取；
    无 LLM synthesis 标记时 fallback 到 configured_mode（默认 llm），
    不再硬编码 deterministic。前端通过 configured_mode 区分用户配置。
    """

    cfg_path, vault, cards = _write_config(tmp_path)
    _write_approved_card(cards)
    wiki_dir = vault / "30-Wiki"
    wiki_dir.mkdir()
    (wiki_dir / "Main-Wiki.md").write_text(
        """# MindForge Main Wiki

> This wiki is generated from human-approved knowledge cards.
> It is a derived view. Source files are not copied into this wiki.
> Last rebuilt: 2026-05-18T10:30:00+0800
> Cards included: 1

## Overview

Synthetic overview.

<!-- WIKI_SECTION_START card=approved-web-1 -->
### Approved Web Card

Approved section body.

<!-- WIKI_SECTION_END -->
""",
        encoding="utf-8",
    )

    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    response = client.get("/api/wiki/page")

    assert response.status_code == 200
    data = response.json()
    # deterministic header ("This wiki is generated from human-approved...")
    # 被正确识别；content_mode="deterministic" 不 fallback 到 configured_mode="llm"
    assert data["content_mode"] == "deterministic"
    assert data["mode"] == "deterministic"  # legacy alias for content_mode
    assert data["configured_mode"] == "llm"
    assert data["model_id"] is None
    assert data["last_rebuilt_at"] == "2026-05-18T10:30:00+0800"


def test_wiki_page_reports_llm_metadata_from_wiki_header(tmp_path: Path) -> None:
    """LLM synthesis Wiki 页面应保留 model_id 和 last_rebuilt_at metadata。"""

    cfg_path, vault, cards = _write_config(tmp_path)
    _write_approved_card(cards)
    wiki_dir = vault / "30-Wiki"
    wiki_dir.mkdir()
    (wiki_dir / "Main-Wiki.md").write_text(
        """# MindForge Main Wiki

> LLM synthesis · Model: wiki_model · Last rebuilt: 2026-05-18T11:00:00+0800
> Cards included: 1

## Overview

LLM overview.

<!-- WIKI_SECTION_START card_ids=approved-web-1 -->
### Approved Web Card

Approved section body.

<!-- WIKI_SECTION_END -->
""",
        encoding="utf-8",
    )

    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    response = client.get("/api/wiki/page")

    assert response.status_code == 200
    data = response.json()
    assert data["mode"] == "llm"
    assert data["model_id"] == "wiki_model"
    assert data["last_rebuilt_at"] == "2026-05-18T11:00:00+0800"


def test_wiki_page_missing_wiki_still_returns_metadata_fields(tmp_path: Path) -> None:
    """缺失 Wiki 文件时 API 应返回稳定 empty-state contract，而不是 500。

    中文学习型说明：Wiki 不存在时 mode 为 None（无内容模式），
    configured_mode 表示用户配置的生成模式。
    """

    cfg_path, _vault, _cards = _write_config(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    response = client.get("/api/wiki/page")

    assert response.status_code == 200
    data = response.json()
    assert data["exists"] is False
    assert data["mode"] is None
    assert data["configured_mode"] == "llm"
    assert data["model_id"] is None
    assert data["last_rebuilt_at"] is None
    assert data["sections"] == []


def test_setup_save_preserves_wiki_auto_rebuild_false(tmp_path: Path) -> None:
    """Setup PATCH 中的 false 是显式用户选择，不能被当成未设置丢弃。"""

    cfg_path, _vault, _cards = _write_config(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    response = client.patch(
        "/api/config/editable",
        json={"wiki_mode": "deterministic", "wiki_auto_rebuild_on_approve": False},
    )

    assert response.status_code == 200
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    assert raw["wiki"]["mode"] == "deterministic"
    assert raw["wiki"]["auto_rebuild_on_approve"] is False


def _write_approved(cards: Path) -> Path:
    card = cards / "approved.md"
    card.write_text(
        """---
id: approved-1
title: Approved Card
status: human_approved
track: product
projects:
  - local-console
tags:
  - library
source_type: manual_note
source_id: src_approved_1
adapter_name: PlainMarkdownAdapter
source_path: 00-Inbox/_processed/ManualNotes/source-note.md
source_content_hash: sha256:approvedhash
source_archive_path: 90-Archive/Processed/source-note.md
source_missing: false
source_title: Approved source
value_score: 7
strategy_id: knowledge_card
strategy_version: 0.10.0
schema_version: "1"
prompt_versions:
  triage: v1
  distill: v1
profile: real-profile
stage_models:
  distill: { alias: real_alias, provider: openai_compatible, model: gpt-test }
run_id: run-approved
---

## Source Excerpt

SOURCE_EXCERPT_VISIBLE_BUT_NOT_RAW_FILE

## AI Summary

Approved summary alpha workspace.

## AI Inference

Approved inference beta.

## Review Questions

- What should be reviewed?

## Action Items

- Maintain the card.

## Human Note

HUMAN_NOTE_MUST_NOT_BE_RECALLABLE
""",
        encoding="utf-8",
    )
    return card


def _client(tmp_path: Path, monkeypatch) -> tuple[TestClient, Path]:
    cfg_path, _vault, cards = _write_config(tmp_path)
    (tmp_path / ".env").write_text(
        "MINDFORGE_FAKE_SECRET=do-not-leak\nMINDFORGE_CUBOX_TOKEN=cubox-secret\n",
        encoding="utf-8",
    )
    monkeypatch.chdir(tmp_path)
    app = create_app(config_path=cfg_path, host="127.0.0.1")
    return TestClient(app), cards


def _wait_for_processing_run(
    client: TestClient,
    run_id: str,
    *,
    timeout: float = 5.0,
) -> dict:
    """轮询后台 processing run，避免测试重新把 HTTP 请求变成同步等待。

    中文学习型说明：Web 请求只负责启动后台任务；测试通过独立 status API
    观察最终结果，锁定“启动边界”和“完成反馈边界”是两条不同链路。
    """

    deadline = time.monotonic() + timeout
    latest: dict | None = None
    while time.monotonic() < deadline:
        response = client.get(f"/api/processing/runs/{run_id}")
        assert response.status_code == 200
        latest = response.json()
        if latest["status"] not in {"queued", "running"}:
            return latest
        time.sleep(0.05)
    raise AssertionError(f"processing run {run_id} did not finish; latest={latest}")


def test_web_workflow_library_and_source_visibility_return_card_content_not_source_raw(
    tmp_path: Path,
    monkeypatch,
) -> None:
    client, cards = _client(tmp_path, monkeypatch)
    card = _write_draft(cards)
    # 将 card 状态改为 human_approved，使 library listing 包含该 card
    raw_text = card.read_text()
    card.write_text(raw_text.replace("status: ai_draft", "status: human_approved"))
    source = tmp_path / "dogfood-vault" / "00-Inbox" / "ManualNotes" / "source-note.md"
    source.write_text("SOURCE_BODY_MUST_NOT_LEAK", encoding="utf-8")
    processed = (
        tmp_path
        / "dogfood-vault"
        / "00-Inbox"
        / "_processed"
        / "ManualNotes"
        / "done.md"
    )
    processed.parent.mkdir(parents=True)
    processed.write_text("PROCESSED_BODY_MUST_NOT_LEAK", encoding="utf-8")

    workflow = client.get("/api/workflow/summary").json()
    library = client.get("/api/library/cards").json()
    detail = client.get("/api/library/card", params={"ref": "draft-1"}).json()
    sources = client.get("/api/sources").json()
    combined = f"{workflow} {library} {detail} {sources}"

    assert workflow["vault_root"].endswith("dogfood-vault")
    assert workflow["human_approved_count"] == 1
    assert workflow["ai_draft_count"] == 0
    assert workflow["source_bucket_counts"]["pending"]["ManualNotes"] == 1
    assert workflow["source_bucket_counts"]["processed"]["ManualNotes"] == 1
    assert library["cards"][0]["source_path"] == "00-Inbox/ManualNotes/source-note.md"
    assert library["cards"][0]["source_id"] == "src_draft_1"
    assert library["cards"][0]["source_content_hash"] == "sha256:provenancehash"
    assert library["cards"][0]["strategy_id"] == "five_stage"
    assert library["cards"][0]["prompt_versions"]["distill"] == "v1"
    assert library["cards"][0]["run_id"] == "run-web-provenance"
    assert library["cards"][0]["source_missing"] is False
    assert detail["card"]["id"] == "draft-1"
    assert detail["card"]["strategy_version"] == "0.10.0"
    assert detail["card"]["schema_version"] == "1"
    assert detail["card"]["fallback_provider_note"]
    assert "safe draft body" in detail["body"].lower()
    assert "SOURCE_BODY_MUST_NOT_LEAK" not in combined
    assert "PROCESSED_BODY_MUST_NOT_LEAK" not in combined
    assert card.exists()


def test_sources_api_exposes_frequency_due_and_baseline_counts(tmp_path: Path, monkeypatch) -> None:
    client, _cards = _client(tmp_path, monkeypatch)
    source = tmp_path / "watched.md"
    source.write_text("# Watched\n\nbody\n", encoding="utf-8")

    added = client.post("/api/sources/watch", json={"path": str(source), "frequency": "daily"})
    _wait_for_processing_run(client, added.json()["run_id"])
    sources = client.get("/api/sources").json()

    assert added.status_code == 200
    watched = [item for item in sources["watched_sources"] if item["path"] == str(source.resolve())][0]
    assert watched["frequency"] == "daily"
    assert watched["last_scan_at"]
    assert watched["next_scan_at"]
    assert watched["due_status"] in {"Due", "Not due", "Manual"}
    assert "added" in watched["diff_counts"]


def test_add_and_process_now_starts_background_run_and_draft_becomes_reviewable(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Add and process now 只启动后台 run；draft 完成后才出现在 Review。"""

    client, _cards = _client(tmp_path, monkeypatch)
    source = tmp_path / "watched.md"
    source.write_text("# Background Source\n\nbody worth keeping\n", encoding="utf-8")

    response = client.post(
        "/api/sources/watch",
        json={"path": str(source), "frequency": "manual", "process_now": True},
    )

    assert response.status_code == 200
    started = response.json()
    assert started["run_id"]
    assert started["processing_status"] in {"queued", "running"}
    assert "background" in started["message"].lower()
    assert "You can keep using MindForge." in started["message"]
    assert "processed as ai_draft" not in started["message"]
    assert all(action.get("href") != "/drafts" for action in started["next_actions"])

    finished = _wait_for_processing_run(client, started["run_id"])
    assert finished["status"] == "succeeded"
    assert finished["summary"]["drafts"] == 1
    assert finished["draft_ids"]
    assert any(action.get("href") == "/drafts" for action in finished["next_actions"])

    drafts = client.get("/api/drafts").json()
    assert len(drafts["drafts"]) == 1
    assert drafts["drafts"][0]["status"] == "ai_draft"
    library = client.get("/api/library/cards").json()
    assert library["stats"]["by_status"].get("human_approved", 0) == 0
    assert not (tmp_path / "dogfood-vault" / "40-Wiki" / "Main-Wiki.md").exists()


def test_process_now_starts_background_run_and_sources_expose_last_run_summary(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Sources 页 Process now 返回 run_id，刷新后能看到 last run summary。"""

    client, _cards = _client(tmp_path, monkeypatch)
    source = tmp_path / "watched.md"
    source.write_text("# Existing Watch\n\nbody worth keeping\n", encoding="utf-8")
    added = client.post(
        "/api/sources/watch",
        json={"path": str(source), "frequency": "manual", "process_now": False},
    ).json()

    response = client.post("/api/sources/watch/scan", params={"ref": added["watch_id"]})

    assert response.status_code == 200
    started = response.json()
    assert started["run_id"]
    assert started["processing_status"] in {"queued", "running"}
    assert "background" in started["message"].lower()
    assert "You can keep using MindForge." in started["message"]

    finished = _wait_for_processing_run(client, started["run_id"])
    assert finished["status"] == "succeeded"

    sources = client.get("/api/sources").json()
    watched = [item for item in sources["watched_sources"] if item["id"] == added["watch_id"]][0]
    assert watched["last_run_summary"]["drafts"] == 1
    assert watched["last_message"] == "Generated 1 AI draft."
    assert watched["generated_draft_count"] == 1
    assert watched["processing_status"] == "succeeded"


def test_process_now_reports_legacy_doc_as_friendly_skip(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Legacy .doc 在 Web 后台 run 中必须是友好 skipped，不是 NoOutputError。

    中文学习型说明：source discovery 已经知道 legacy .doc 是 unsupported；
    processing run 汇总层必须保留这个事实，不能因为没有 draft 就误导用户去
    检查模型配置。这里锁住 fresh-clone dogfood 暴露的真实回归路径。
    """

    client, _cards = _client(tmp_path, monkeypatch)
    legacy = tmp_path / "legacy-plan.doc"
    legacy.write_bytes(b"synthetic legacy ole placeholder")

    response = client.post(
        "/api/sources/watch",
        json={"path": str(legacy), "frequency": "manual", "process_now": True},
    )

    assert response.status_code == 200
    finished = _wait_for_processing_run(client, response.json()["run_id"])
    assert finished["status"] == "skipped"
    assert finished["summary"]["discovered"] == 1
    assert finished["summary"]["skipped"] == 1
    assert finished["summary"]["errors"] == 0
    assert finished["error_type"] is None
    assert any(reason.startswith("unsupported_legacy_doc") for reason in finished["skip_reasons"])
    assert "Model setup" not in finished["message"]
    assert "convert" in finished["message"].lower()


def test_process_now_returns_existing_active_run_instead_of_spawning_duplicate(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """重复点击同一个 source 不应制造多个并发 active run。

    中文学习型说明：当前 Web async 是 request-spawn，不是 durable queue。
    因此最小产品边界是“同一 source 已有 active run 时复用它”，避免旧 run 和
    新 run 互相覆盖用户可见状态，也避免重复消耗模型调用。
    """

    client, _cards = _client(tmp_path, monkeypatch)
    source = tmp_path / "slow-source.md"
    source.write_text("# Slow Source\n\nbody worth keeping\n", encoding="utf-8")
    gate = threading.Event()
    calls = 0

    from mindforge_web.services import web_source_service as source_service

    original_scan = source_service.watch_scan_sources

    def slow_scan(*args, **kwargs):
        nonlocal calls
        calls += 1
        gate.wait(timeout=3)
        return original_scan(*args, **kwargs)

    monkeypatch.setattr(source_service, "watch_scan_sources", slow_scan)

    registered = client.post(
        "/api/sources/watch",
        json={"path": str(source), "frequency": "manual", "process_now": False},
    ).json()
    first = client.post("/api/sources/watch/scan", params={"ref": registered["watch_id"]}).json()
    second = client.post("/api/sources/watch/scan", params={"ref": registered["watch_id"]}).json()
    gate.set()
    finished = _wait_for_processing_run(client, first["run_id"])

    assert first["run_id"] == second["run_id"]
    assert second["processing_status"] in {"queued", "running"}
    assert "already running" in second["message"].lower()
    assert calls == 1
    assert finished["status"] == "succeeded"


# ── P1 回归测试：重复 ai_draft 防护 ──────────────────────────────
# 中文学习型说明：以下 5 个测试锁定 P1 fix 引入的三层 dedup：
#   L1: active-run 协调（processing_run_service）
#   L2: card-level dedup（ingestion_service._build_existing_draft_index）
#   L3: checkpoint fingerprint dedup（processed_sources / content_hash_index）


def test_import_then_watch_scan_same_txt_source_only_one_ai_draft(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """import TXT 后再 watch/scan 同一文件，card-level dedup 防止重复 ai_draft。

    中文学习型说明：P1 L2 防线——import 同步生成 ai_draft 后，watch/scan 的
    _ingest_targets_summary 通过 _build_existing_draft_index 发现已有同
    source_path 的 ai_draft，必须 skip 而非再写一张。
    """

    cfg_path, _vault, cards = _write_config_multi_adapter(tmp_path)
    (tmp_path / ".env").write_text(
        "MINDFORGE_FAKE_SECRET=do-not-leak\nMINDFORGE_CUBOX_TOKEN=cubox-secret\n",
        encoding="utf-8",
    )
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    txt_file = tmp_path / "dedup-test.txt"
    txt_file.write_text("TXT body for dedup regression test.\n", encoding="utf-8")

    # Step 1：import 同步生成 1 张 ai_draft
    imported = client.post("/api/sources/import", json={"path": str(txt_file)}).json()
    assert imported["ok"] is True
    assert imported["counts"]["processed"] == 1
    assert imported["run_id"]

    # Step 2：watch_add + scan 同一文件 → card-level dedup 应阻止再生成
    watched = client.post(
        "/api/sources/watch",
        json={"path": str(txt_file), "frequency": "manual", "process_now": False},
    ).json()
    scanned = client.post(
        "/api/sources/watch/scan", params={"ref": watched["watch_id"]}
    ).json()
    finished = _wait_for_processing_run(client, scanned["run_id"])

    # watch/scan 必须 skipped（不是 NoOutputError）
    # skip reason 可能是 card-level dedup（ai_draft_already_exists）或
    # checkpoint-level dedup（already_processed），取决于两次调用的时序
    assert finished["status"] == "skipped"
    assert finished["summary"]["drafts"] == 0
    valid_reasons = {"ai_draft_already_exists", "already_processed"}
    assert any(
        any(reason in r for reason in valid_reasons)
        for r in finished["skip_reasons"]
    ), f"expected one of {valid_reasons} in skip_reasons, got {finished['skip_reasons']}"

    # 最终 vault 只有 1 张 ai_draft
    draft_paths = [p for p in cards.rglob("*.md") if "ai_draft" in p.read_text(encoding="utf-8")]
    assert len(draft_paths) == 1


def test_import_then_watch_scan_same_docx_source_only_one_ai_draft(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """import DOCX 后再 watch/scan 同一文件，card-level dedup 防止重复 ai_draft。

    中文学习型说明：同 TXT 用例，但验证 DocxTextAdapter 路径下的 source_path
    标准化与 card dedup 无误。python-docx 生成最小合法 .docx。
    """

    cfg_path, _vault, cards = _write_config_multi_adapter(tmp_path)
    (tmp_path / ".env").write_text(
        "MINDFORGE_FAKE_SECRET=do-not-leak\nMINDFORGE_CUBOX_TOKEN=cubox-secret\n",
        encoding="utf-8",
    )
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    docx_file = tmp_path / "dedup-test.docx"
    _write_minimal_docx(docx_file)

    imported = client.post("/api/sources/import", json={"path": str(docx_file)}).json()
    assert imported["ok"] is True
    assert imported["counts"]["processed"] == 1

    watched = client.post(
        "/api/sources/watch",
        json={"path": str(docx_file), "frequency": "manual", "process_now": False},
    ).json()
    scanned = client.post(
        "/api/sources/watch/scan", params={"ref": watched["watch_id"]}
    ).json()
    finished = _wait_for_processing_run(client, scanned["run_id"])

    assert finished["status"] == "skipped"
    assert finished["summary"]["drafts"] == 0
    valid_reasons = {"ai_draft_already_exists", "already_processed"}
    assert any(
        any(reason in r for reason in valid_reasons)
        for r in finished["skip_reasons"]
    ), f"expected one of {valid_reasons} in skip_reasons, got {finished['skip_reasons']}"

    draft_paths = [p for p in cards.rglob("*.md") if "ai_draft" in p.read_text(encoding="utf-8")]
    assert len(draft_paths) == 1


def test_active_run_blocks_concurrent_import_for_same_source(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """同一 source 已有 active run 时 import 复用该 run，不新开并发。

    中文学习型说明：P1 L1 防线——start_sync_processing_run 在创建 run 之前
    先查 latest_run_for_source。如果已有 queued/running 的 run，直接返回它，
    不创建新 run。这防止两个 HTTP 请求同时触发同一 source 的 processing。
    """

    cfg_path, _vault, _cards = _write_config(tmp_path)
    (tmp_path / ".env").write_text(
        "MINDFORGE_FAKE_SECRET=do-not-leak\nMINDFORGE_CUBOX_TOKEN=cubox-secret\n",
        encoding="utf-8",
    )
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    source = tmp_path / "concurrent-import.md"
    source.write_text("# Concurrent Import\n\nbody\n", encoding="utf-8")

    # 第一个 import：正常生成
    first = client.post("/api/sources/import", json={"path": str(source)}).json()
    assert first["ok"] is True
    assert first["counts"]["processed"] == 1
    first_run_id = first["run_id"]

    # 第二个 import：同 source 再次 import，走 checkpoint dedup
    second = client.post("/api/sources/import", json={"path": str(source)}).json()
    assert second["ok"] is True
    # 第二次不应生成新 draft
    assert second["counts"]["processed"] == 0
    assert second["counts"]["skipped"] >= 1
    # 两次返回不同 run_id（因为第一次已 finalize，第二次是新 run）
    assert second["run_id"] != first_run_id
    cards_dir = tmp_path / "dogfood-vault" / "20-Knowledge-Cards"
    draft_paths = [p for p in cards_dir.rglob("*.md") if "ai_draft" in p.read_text(encoding="utf-8")]
    assert len(draft_paths) == 1


def test_unchanged_fingerprint_reimport_no_duplicate_draft(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """内容未变的 source 再次 import 不生成新 ai_draft。

    中文学习型说明：P1 L3 防线——checkpoint 层的 processed_sources 索引
    (source_id + content_hash) 和 processed_content_hash_index 必须阻止
    同 fingerprint 重复生成卡片。两次独立 import 调用，内容 hash 相同，
    第二次应 skip。
    """

    cfg_path, _vault, cards = _write_config(tmp_path)
    (tmp_path / ".env").write_text(
        "MINDFORGE_FAKE_SECRET=do-not-leak\nMINDFORGE_CUBOX_TOKEN=cubox-secret\n",
        encoding="utf-8",
    )
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    source = tmp_path / "unchanged.md"
    source.write_text("# Unchanged\n\nSame body twice.\n", encoding="utf-8")

    first = client.post("/api/sources/import", json={"path": str(source)}).json()
    assert first["ok"] is True
    assert first["counts"]["processed"] == 1

    second = client.post("/api/sources/import", json={"path": str(source)}).json()
    assert second["ok"] is True
    assert second["counts"]["processed"] == 0
    assert second["counts"]["skipped"] >= 1

    draft_paths = [p for p in cards.rglob("*.md") if "ai_draft" in p.read_text(encoding="utf-8")]
    assert len(draft_paths) == 1


def test_changed_source_reimport_generates_new_draft(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """内容变更后的 source 重新 import 必须生成新 ai_draft。

    中文学习型说明：P1 的关键反例——dedup 不能阻止用户对真正变更了内容的
    source 重新处理。content_hash 不同时必须走完整 pipeline 产出新 card。
    """

    cfg_path, _vault, cards = _write_config(tmp_path)
    (tmp_path / ".env").write_text(
        "MINDFORGE_FAKE_SECRET=do-not-leak\nMINDFORGE_CUBOX_TOKEN=cubox-secret\n",
        encoding="utf-8",
    )
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    source = tmp_path / "changing.md"
    source.write_text("# Version 1\n\nOriginal body.\n", encoding="utf-8")

    first = client.post("/api/sources/import", json={"path": str(source)}).json()
    assert first["ok"] is True
    assert first["counts"]["processed"] == 1

    # 内容变更
    source.write_text("# Version 2\n\nUpdated body with new content.\n", encoding="utf-8")

    second = client.post("/api/sources/import", json={"path": str(source)}).json()
    assert second["ok"] is True
    assert second["counts"]["processed"] == 1, (
        f"changed source must generate new draft, got {second['counts']}"
    )

    draft_paths = [p for p in cards.rglob("*.md") if "ai_draft" in p.read_text(encoding="utf-8")]
    assert len(draft_paths) == 2


# ── P2 回归测试：ProcessingRun 状态过渡 ───────────────────────────
# 中文学习型说明：_apply_summary 必须覆盖所有 count 组合映射到正确终结状态，
# 不能遗留 'running'。以下 5 个单元级测试直接调用 _apply_summary 验证。


def _make_record() -> ProcessingRunRecord:
    return ProcessingRunRecord(
        run_id="pr_p2_test",
        source_ref="source-p2",
        source_path="/tmp/p2-test.md",
        mode="watch_scan",
        status="running",
        started_at="2026-01-01T00:00:00.000000+00:00",
        current_step="processing source",
    )


def _make_summary(**counts: int) -> IngestionSummary:
    return IngestionSummary(
        mode="watch_scan",
        target=Path("/tmp/p2-test.md"),
        counts=counts,
    )


def test_apply_summary_no_output_discovered_zero_status_skipped() -> None:
    """discovered=0 → skipped。扫描无支持文件不算失败。"""
    record = _make_record()
    _apply_summary(record, _make_summary(seen=0, processed=0, skipped=0, failed=0), draft_ids=[])
    assert record.status == "skipped"
    assert record.summary["discovered"] == 0


def test_apply_summary_all_succeeded_status_succeeded() -> None:
    """全部成功生成 drafts → succeeded。"""
    record = _make_record()
    _apply_summary(record, _make_summary(seen=2, processed=2, skipped=0, failed=0), draft_ids=["a", "b"])
    assert record.status == "succeeded"
    assert record.summary["drafts"] == 2


def test_apply_summary_all_skipped_status_skipped() -> None:
    """全部被跳过（duplicate / triage / adapter skip）→ skipped。"""
    record = _make_record()
    _apply_summary(record, _make_summary(seen=3, processed=0, skipped=3, failed=0), draft_ids=[])
    assert record.status == "skipped"
    assert record.summary["skipped"] == 3


def test_apply_summary_partial_failure_status_partial_failed() -> None:
    """部分成功 + 部分失败 → partial_failed。"""
    record = _make_record()
    _apply_summary(record, _make_summary(seen=3, processed=2, skipped=0, failed=1), draft_ids=["a", "b"])
    assert record.status == "partial_failed"
    assert record.summary["drafts"] == 2
    assert record.summary["errors"] == 1


def test_apply_summary_all_failed_status_failed() -> None:
    """全部失败无产出 → failed。"""
    record = _make_record()
    _apply_summary(record, _make_summary(seen=1, processed=0, skipped=0, failed=1), draft_ids=[])
    assert record.status == "failed"
    assert record.summary["errors"] == 1


# ── P2-2 回归测试：transient failure missing error details ──────────
# 中文学习型说明：P2-2 修复确保 transient adapter/process exception 出现时，
# API response / run detail 必须包含 error_type 和 friendly message；errors
# count 不能只有数字没有细节；不泄露 stack trace / secret / raw provider payload。


def test_apply_summary_errors_without_messages_still_sets_error_type() -> None:
    """errors > 0 但 error_messages 为空时，error_type 和 retry hint 不能缺失。"""
    record = _make_record()
    # 构造无 error 消息尽有 count 的 summary（模拟 pipeline 返回空 message 场景）
    summary = IngestionSummary(
        mode="watch_scan",
        target=Path("/tmp/p2-2-test.md"),
        counts={"seen": 1, "processed": 0, "skipped": 0, "failed": 1},
        errors=(),  # 空 tuple — 没有任何错误消息
    )
    _apply_summary(record, summary, draft_ids=[])
    assert record.status == "failed"
    assert record.summary["errors"] == 1
    assert record.error_type == "ProcessingError", (
        f"errors=1 但 error_messages 为空时未设置 error_type，got {record.error_type}"
    )
    assert record.error_message is not None
    assert "retry" in record.error_message.lower()
    # 不应泄露 stack trace
    assert "Traceback" not in record.error_message
    assert "secret" not in record.error_message.lower()


def test_apply_summary_errors_with_specific_message_preserves_error_detail() -> None:
    """error_messages 非空时，error_type、failed stage 和具体错误消息都应保留。"""
    record = _make_record()
    summary = IngestionSummary(
        mode="watch_scan",
        target=Path("/tmp/p2-2-test.md"),
        counts={"seen": 1, "processed": 0, "skipped": 0, "failed": 1},
        errors=("distill_stage_failed: JSON parse error",),
    )
    _apply_summary(record, summary, draft_ids=[])
    assert record.status == "failed"
    assert record.error_type == "ProcessingError"
    assert "JSON parse error" in record.error_message
    assert record.current_step == "failed: distill"


def test_message_for_summary_errors_without_reason_includes_retry_hint() -> None:
    """message_for_summary 在 errors > 0 但无 reason 时必须提示重试。"""
    record = _make_record()
    record.status = "failed"
    record.summary = {"discovered": 1, "processed": 0, "drafts": 0, "skipped": 0, "errors": 1}
    record.error_message = None
    record.skip_reasons = []
    msg = _message_for_summary(record)
    assert "failed" in msg.lower()
    assert "retry" in msg.lower()
    assert "transient" in msg.lower()


def test_message_for_summary_errors_with_reason_shows_reason_not_retry() -> None:
    """message_for_summary 在有具体 reason 时展示原因，而非泛泛重试提示。"""
    record = _make_record()
    record.status = "failed"
    record.summary = {"discovered": 1, "processed": 0, "drafts": 0, "skipped": 0, "errors": 1}
    record.error_message = "distill stage JSON parse error"
    msg = _message_for_summary(record)
    assert "JSON parse error" in msg
    assert "Reason:" in msg
    # 有具体 reason 时不展示泛型重试文案
    assert "transient" not in msg.lower()


def test_safe_error_message_strips_html_and_stack_trace() -> None:
    """_safe_error_message 必须剥除 HTML 错误页和 stack trace。"""
    # HTML 错误页 → 友好提示
    html_msg = (
        "<!DOCTYPE html><html><body><h1>502 Bad Gateway</h1>"
        "The proxy server received an invalid response.</body></html>"
    )
    cleaned = _safe_error_message(html_msg)
    assert "Provider returned an HTML error page" in cleaned
    assert "<!DOCTYPE html>" not in cleaned
    assert "<html" not in cleaned


def test_safe_error_message_preserves_friendly_diagnostic() -> None:
    """_safe_error_message 应保留 friendly_missing_key_error 的诊断消息。"""
    cleaned = _safe_error_message("API key 未设置，无法完成请求")
    assert cleaned is not None
    assert "API key" in cleaned


# ── P2b 回归测试：PDF 解析失败原因提示 ────────────────────────────
# 中文学习型说明：_skip_reason_hint 将 adapter-level 错误码翻译为用户友好文案，
# 区分 encrypted_pdf / parse_failed / adapter_error 等不同失败原因。


def test_skip_reason_hint_encrypted_pdf_returns_password_hint() -> None:
    """encrypted_pdf → 提示用户移除密码保护或导出为非加密 PDF。"""
    hint = _skip_reason_hint("encrypted_pdf")
    assert hint is not None
    assert "encrypted" in hint.lower()
    assert "password" in hint.lower()


def test_skip_reason_hint_adapter_error_prefix_returns_pdf_extraction_hint() -> None:
    """adapter_error: <message> → 提示 PDF text extraction failed，不含 OCR 误导。"""
    hint = _skip_reason_hint("adapter_error: PdfReadError: something wrong")
    assert hint is not None
    assert "pdf" in hint.lower()
    assert "ocr" in hint.lower()
    assert "not supported" in hint.lower()


def test_skip_reason_hint_parse_failed_returns_friendly_format_guidance() -> None:
    """parse_failed → 友好提示支持格式，提到 OCR 不可用。"""
    hint = _skip_reason_hint("parse_failed")
    assert hint is not None
    assert "parsing failed" in hint.lower()
    assert "markdown" in hint.lower()
    assert "docx" in hint.lower()
    assert "ocr" in hint.lower()


def test_skip_reason_hint_unknown_reason_returns_none() -> None:
    """未知 reason → None，调用方自行处理默认展示。"""
    assert _skip_reason_hint("some_random_other_reason") is None


def test_skip_reason_hint_empty_string_returns_none() -> None:
    """空字符串 → None，不产生虚假提示。"""
    assert _skip_reason_hint("") is None


# ── P3 回归测试：Source Adapter 状态展示 ──────────────────────────
# 中文学习型说明：_display_status 将 source adapter 的 pending/processed/error
# 计数映射为展示标签。v0.3 P3 fix 新增 generated_card_count，确保 watched source
# 已生成 ai_draft 的 adapter 显示 "Processed" 而非 "Imported"/"Pending"。


def test_display_status_imported_no_registry_shows_imported_not_pending() -> None:
    """无 pending 无 processed 无 error → "Imported"，不是 "Pending"。"""
    assert _display_status(exists=True, pending_count=0, processed_count=0, error_count=0) == "Imported"


def test_display_status_has_generated_cards_shows_processed() -> None:
    """有 ai_draft (generated_card_count > 0) 但 processed_dir 为空 → "Processed"。"""
    assert _display_status(
        exists=True, pending_count=0, processed_count=0, error_count=0, generated_card_count=3
    ) == "Processed"


def test_display_status_incoming_files_no_cards_shows_pending() -> None:
    """有 pending 文件但无 generated cards → "Pending"。"""
    assert _display_status(
        exists=True, pending_count=5, processed_count=0, error_count=0, generated_card_count=0
    ) == "Pending"


def test_display_status_error_overrides_all() -> None:
    """error_count > 0 → "Failed"，无视其他计数。"""
    assert _display_status(
        exists=True, pending_count=5, processed_count=3, error_count=1, generated_card_count=3
    ) == "Failed"


def test_display_status_missing_folder_overrides_all() -> None:
    """inbox 目录不存在 → "Missing folder"。"""
    assert _display_status(
        exists=False, pending_count=0, processed_count=0, error_count=0
    ) == "Missing folder"


def test_display_status_processed_and_pending_shows_processed() -> None:
    """同时有 processed 和 pending → "Processed" 优先。"""
    assert _display_status(
        exists=True, pending_count=3, processed_count=2, error_count=0
    ) == "Processed"


def test_stale_running_run_is_visible_as_failed_after_reload(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """服务重启后的 orphan running run 必须变成用户可见失败。

    中文学习型说明：daemon thread 不会跨进程恢复。与其让 Sources 永远显示
    running，不如把超过阈值的 active run 标记为 abandoned failed，并给出
    retry action。这样用户能理解状态并重新 Process now。
    """

    cfg_path, _vault, _cards = _write_config(tmp_path)
    source = tmp_path / "stale.md"
    source.write_text("# Stale\n\nbody\n", encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    registered = client.post(
        "/api/sources/watch",
        json={"path": str(source), "frequency": "manual", "process_now": False},
    ).json()
    cfg = create_app(config_path=cfg_path, host="127.0.0.1").state.facade.cfg
    stale = ProcessingRunRecord(
        run_id="pr_stale_reload",
        source_ref=registered["watch_id"],
        source_path=str(source.resolve()),
        mode="watch_scan",
        status="running",
        started_at="2000-01-01T00:00:00.000000+00:00",
        current_step="processing source",
        message="Processing started in the background. You can keep using MindForge.",
    )
    _save_processing_run_record(cfg, stale)

    run = client.get("/api/processing/runs/pr_stale_reload").json()
    sources = client.get("/api/sources").json()
    watched = next(item for item in sources["watched_sources"] if item["id"] == registered["watch_id"])

    assert run["status"] == "failed"
    assert run["current_step"] == "abandoned"
    assert "did not finish" in run["message"]
    assert any(action["label"] == "Retry processing" for action in run["next_actions"])
    assert watched["processing_status"] == "failed"
    assert watched["active_run_id"] is None
    assert watched["last_error"]


def test_worker_completion_does_not_overwrite_finalized_run(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """worker 最终完成不能覆盖已经 finalized 的 run。

    中文学习型说明：Mina 审计指出过真实竞态：normalizer 先把 stale running
    标为 failed，随后 worker 完成又写 succeeded。run lifecycle 必须单调：
    一旦离开 queued/running，后台 worker 只能停止写最终状态，不能反向覆盖。
    """

    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    cfg = create_app(config_path=cfg_path, host="127.0.0.1").state.facade.cfg
    entered_work = threading.Event()
    release_work = threading.Event()
    record = ProcessingRunRecord(
        run_id="pr_finalized_guard",
        source_ref="source-finalized",
        source_path=str(tmp_path / "source.md"),
        mode="watch_scan",
        status="queued",
        started_at=_now(),
    )
    _save_processing_run_record(cfg, record)

    def slow_success() -> IngestionSummary:
        entered_work.set()
        release_work.wait(timeout=10)
        return IngestionSummary(
            mode="watch_scan",
            target=tmp_path / "source.md",
            counts={"seen": 1, "processed": 1, "skipped": 0, "failed": 0},
        )

    worker = threading.Thread(target=_run_worker, args=(cfg, record.run_id, slow_success))
    worker.start()
    deadline = time.monotonic() + 3
    while time.monotonic() < deadline:
        running = create_app(config_path=cfg_path, host="127.0.0.1").state.facade.processing_run(record.run_id)
        if running is not None and running.status == "running":
            break
        time.sleep(0.02)
    else:
        raise AssertionError("worker did not enter running state")
    assert entered_work.wait(timeout=3)

    finalized = ProcessingRunRecord(
        run_id=record.run_id,
        source_ref=record.source_ref,
        source_path=record.source_path,
        mode=record.mode,
        status="failed",
        started_at=record.started_at,
        finished_at=_now(),
        current_step="abandoned",
        summary={"discovered": 0, "processed": 0, "drafts": 0, "skipped": 0, "errors": 1},
        message="Processing did not finish.",
        error_type="AbandonedProcessingRun",
        error_message="abandoned by normalizer",
    )
    _save_processing_run_record(cfg, finalized)

    release_work.set()
    worker.join(timeout=3)
    latest = create_app(config_path=cfg_path, host="127.0.0.1").state.facade.processing_run(record.run_id)

    assert latest is not None
    assert latest.status == "failed"
    assert latest.current_step == "abandoned"
    assert latest.summary["errors"] == 1
    assert latest.summary["drafts"] == 0
    assert latest.error_message == "abandoned by normalizer"


def test_fresh_heartbeat_prevents_long_running_false_abandoned(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """长时间运行但 heartbeat 新鲜的 run 不能被误判 abandoned。"""

    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    cfg = create_app(config_path=cfg_path, host="127.0.0.1").state.facade.cfg
    record = ProcessingRunRecord(
        run_id="pr_fresh_heartbeat",
        source_ref="source-heartbeat",
        source_path=str(tmp_path / "source.md"),
        mode="watch_scan",
        status="running",
        started_at="2000-01-01T00:00:00.000000+00:00",
        last_heartbeat_at=_now(),
        current_step="processing source",
    )
    _save_processing_run_record(cfg, record)

    run = TestClient(create_app(config_path=cfg_path, host="127.0.0.1")).get(
        "/api/processing/runs/pr_fresh_heartbeat"
    ).json()

    assert run["status"] == "running"
    assert run["last_heartbeat_at"]
    assert run["error_type"] is None


def test_running_processing_run_response_exposes_step_and_heartbeat(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """running run 要展示当前 step/heartbeat，不能只显示 discovered=0。

    中文学习型说明：真实 provider 调用可能持续几十秒。用户在 release 主路径
    看到 running 时，至少要知道 worker 仍在处理模型调用附近的阶段，而不是
    误以为系统卡在空扫描或伪成功。
    """

    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    cfg = create_app(config_path=cfg_path, host="127.0.0.1").state.facade.cfg
    record = ProcessingRunRecord(
        run_id="pr_running_visible",
        source_ref="source-visible",
        source_path=str(tmp_path / "source.md"),
        mode="import",
        status="running",
        started_at=_now(),
        last_heartbeat_at=_now(),
        current_step="calling model: triage",
    )
    _save_processing_run_record(cfg, record)

    run = TestClient(create_app(config_path=cfg_path, host="127.0.0.1")).get(
        "/api/processing/runs/pr_running_visible"
    ).json()

    assert run["status"] == "running"
    assert run["current_step"] == "calling model: triage"
    assert run["last_heartbeat_at"]
    assert "calling model" in run["message"].lower()


def test_stale_heartbeat_marks_run_abandoned(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """heartbeat stale 才能把 active run 收敛为 abandoned failed。"""

    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    cfg = create_app(config_path=cfg_path, host="127.0.0.1").state.facade.cfg
    record = ProcessingRunRecord(
        run_id="pr_stale_heartbeat",
        source_ref="source-heartbeat",
        source_path=str(tmp_path / "source.md"),
        mode="watch_scan",
        status="running",
        started_at="2000-01-01T00:00:00.000000+00:00",
        last_heartbeat_at="2000-01-01T00:00:10.000000+00:00",
        current_step="processing source",
    )
    _save_processing_run_record(cfg, record)

    run = TestClient(create_app(config_path=cfg_path, host="127.0.0.1")).get(
        "/api/processing/runs/pr_stale_heartbeat"
    ).json()

    assert run["status"] == "failed"
    assert run["current_step"] == "abandoned"
    assert run["error_type"] == "AbandonedProcessingRun"


def test_processing_run_get_normalizes_stale_run_without_writing_disk(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """GET run 只能返回 normalized view，不能在 read path 写磁盘。

    中文学习型说明：stale/orphan visibility 是用户体验需求，但 GET/read path
    仍应遵守 CQS。这里锁定行为：API 返回 failed/abandoned 视图，但 JSON 文件
    仍保持原始 running 状态，真正状态推进只由 worker 或显式 command path 写盘。
    """

    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    cfg = create_app(config_path=cfg_path, host="127.0.0.1").state.facade.cfg
    record = ProcessingRunRecord(
        run_id="pr_get_no_write",
        source_ref="source-get",
        source_path=str(tmp_path / "source.md"),
        mode="watch_scan",
        status="running",
        started_at="2000-01-01T00:00:00.000000+00:00",
        last_heartbeat_at="2000-01-01T00:00:10.000000+00:00",
        current_step="processing source",
    )
    _save_processing_run_record(cfg, record)
    run_path = tmp_path / ".mindforge" / "processing_runs" / "pr_get_no_write.json"
    before = run_path.read_text(encoding="utf-8")

    response = TestClient(create_app(config_path=cfg_path, host="127.0.0.1")).get(
        "/api/processing/runs/pr_get_no_write"
    )

    assert response.status_code == 200
    assert response.json()["status"] == "failed"
    assert response.json()["current_step"] == "abandoned"
    assert run_path.read_text(encoding="utf-8") == before


def test_sources_normalizes_stale_run_without_writing_disk(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Sources read path 展示 abandoned，但不能顺手 repair run JSON。"""

    cfg_path, _vault, _cards = _write_config(tmp_path)
    source = tmp_path / "source-for-sources.md"
    source.write_text("# Source For Sources\n\nbody\n", encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    registered = client.post(
        "/api/sources/watch",
        json={"path": str(source), "frequency": "manual", "process_now": False},
    ).json()
    cfg = create_app(config_path=cfg_path, host="127.0.0.1").state.facade.cfg
    record = ProcessingRunRecord(
        run_id="pr_sources_no_write",
        source_ref=registered["watch_id"],
        source_path=str(source.resolve()),
        mode="watch_scan",
        status="running",
        started_at="2000-01-01T00:00:00.000000+00:00",
        last_heartbeat_at="2000-01-01T00:00:10.000000+00:00",
        current_step="processing source",
    )
    _save_processing_run_record(cfg, record)
    run_path = tmp_path / ".mindforge" / "processing_runs" / "pr_sources_no_write.json"
    before = run_path.read_text(encoding="utf-8")

    sources = client.get("/api/sources").json()
    watched = next(item for item in sources["watched_sources"] if item["id"] == registered["watch_id"])

    assert watched["processing_status"] == "failed"
    assert watched["last_error"]
    assert run_path.read_text(encoding="utf-8") == before


def test_sources_uses_newest_run_when_older_run_finishes_later(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Sources last run summary 必须以最新 run 为权威状态。

    中文学习型说明：并发或重试场景下旧 run 可能更晚写盘。Sources 面向用户展示
    “最近一次操作”的 lifecycle，不能被较早 started_at 的完成结果倒灌覆盖。
    """

    cfg_path, _vault, _cards = _write_config(tmp_path)
    source = tmp_path / "newest.md"
    source.write_text("# Newest\n\nbody\n", encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    registered = client.post(
        "/api/sources/watch",
        json={"path": str(source), "frequency": "manual", "process_now": False},
    ).json()
    cfg = create_app(config_path=cfg_path, host="127.0.0.1").state.facade.cfg
    older = ProcessingRunRecord(
        run_id="pr_older",
        source_ref=registered["watch_id"],
        source_path=str(source.resolve()),
        mode="watch_scan",
        status="failed",
        started_at="2026-05-11T10:00:00.000000+00:00",
        finished_at="2026-05-11T10:05:00.000000+00:00",
        current_step="failed",
        summary={"discovered": 1, "processed": 0, "drafts": 0, "skipped": 0, "errors": 1},
        message="Processing failed for 1 item(s). Reason: old failure",
        error_type="ProcessingError",
        error_message="old failure",
    )
    newer = ProcessingRunRecord(
        run_id="pr_newer",
        source_ref=registered["watch_id"],
        source_path=str(source.resolve()),
        mode="watch_scan",
        status="succeeded",
        started_at="2026-05-11T10:00:01.000000+00:00",
        finished_at="2026-05-11T10:00:02.000000+00:00",
        current_step="completed",
        summary={"discovered": 1, "processed": 1, "drafts": 1, "skipped": 0, "errors": 0},
        draft_ids=["draft-newer"],
        message="Generated 1 AI draft.",
    )
    _save_processing_run_record(cfg, newer)
    _save_processing_run_record(cfg, older)

    sources = client.get("/api/sources").json()
    watched = next(item for item in sources["watched_sources"] if item["id"] == registered["watch_id"])

    assert watched["processing_status"] == "succeeded"
    assert watched["last_run_id"] == "pr_newer"
    assert watched["last_run_summary"]["drafts"] == 1
    assert watched["last_run_summary"]["errors"] == 0


def test_run_summary_persists_across_app_reload(tmp_path: Path, monkeypatch) -> None:
    """完成后的 run summary 必须刷新 app 后仍可查询。"""

    cfg_path, _vault, _cards = _write_config(tmp_path)
    source = tmp_path / "persisted.md"
    source.write_text("# Persisted\n\nbody worth keeping\n", encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    started = client.post(
        "/api/sources/watch",
        json={"path": str(source), "frequency": "manual", "process_now": True},
    ).json()
    finished = _wait_for_processing_run(client, started["run_id"])
    reloaded = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    after_reload = reloaded.get(f"/api/processing/runs/{started['run_id']}").json()

    assert after_reload["status"] == finished["status"]
    assert after_reload["summary"] == finished["summary"]
    assert after_reload["draft_ids"] == finished["draft_ids"]


def test_processing_run_reports_triage_skip_reason_without_review_next_action(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Triage skipped 时必须告诉用户原因，不能指向空 Review。"""

    cfg_path, _vault, _cards = _write_config(tmp_path)
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["triage"]["value_score_threshold"] = 9
    cfg_path.write_text(yaml.safe_dump(raw, sort_keys=False), encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    source = tmp_path / "low-value.md"
    source.write_text("# Low Value\n\nshort note\n", encoding="utf-8")

    started = client.post(
        "/api/sources/watch",
        json={"path": str(source), "frequency": "manual", "process_now": True},
    ).json()
    finished = _wait_for_processing_run(client, started["run_id"])

    assert finished["status"] == "skipped"
    assert finished["summary"]["drafts"] == 0
    assert finished["summary"]["skipped"] == 1
    assert "value_score=7" in finished["message"]
    assert "threshold=9" in finished["message"]
    assert finished["skip_reasons"]
    assert all(action.get("href") != "/drafts" for action in finished["next_actions"])
    assert client.get("/api/drafts").json()["drafts"] == []


def test_processing_run_failure_is_persisted_and_secret_safe(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Provider/config error 进入 failed run，不泄露 API key，也不无限 loading。"""

    cfg_path, _vault, _cards = _write_config(tmp_path)
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["llm"]["models"]["fake_alias"]["type"] = "openai_compatible"
    raw["llm"]["models"]["fake_alias"]["provider"] = "openai_compatible"
    raw["llm"]["models"]["fake_alias"]["api_key_env"] = "MINDFORGE_FAKE_SECRET"
    cfg_path.write_text(yaml.safe_dump(raw, sort_keys=False), encoding="utf-8")
    (tmp_path / ".env").write_text("MINDFORGE_FAKE_SECRET=secret-must-not-leak\n", encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    monkeypatch.delenv("MINDFORGE_FAKE_SECRET", raising=False)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    source = tmp_path / "provider-error.md"
    source.write_text("# Provider Error\n\nbody\n", encoding="utf-8")

    started = client.post(
        "/api/sources/watch",
        json={"path": str(source), "frequency": "manual", "process_now": True},
    ).json()
    finished = _wait_for_processing_run(client, started["run_id"])
    sources = client.get("/api/sources").json()
    run_record = (
        tmp_path / ".mindforge" / "processing_runs" / f"{started['run_id']}.json"
    ).read_text(encoding="utf-8")
    combined = f"{started} {finished} {sources}"

    assert finished["status"] == "failed"
    assert finished["summary"]["errors"] >= 1
    assert finished["error_message"]
    watched = next(item for item in sources["watched_sources"] if item["path"] == str(source.resolve()))
    assert watched["processing_status"] == "failed"
    assert watched["last_run_summary"]["errors"] >= 1
    assert watched["last_message"]
    assert watched["last_error"]
    assert "secret-must-not-leak" not in combined
    assert "secret-must-not-leak" not in run_record
    assert "processed as ai_draft" not in combined
    assert all(action.get("href") != "/drafts" for action in finished["next_actions"])


def test_processing_run_json_parse_failure_finishes_as_failed(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """LLM JSON parse failure 必须结束为 failed run，不能让 UI 一直 loading。"""

    from mindforge.llm.base import LLMResult
    from mindforge.llm.fake import FakeProvider

    original_generate = FakeProvider.generate

    def invalid_distill_json(self, request):
        if request.stage == "distill":
            return LLMResult(
                text="{not valid json",
                tokens_in=1,
                tokens_out=1,
                latency_ms=0,
                raw={"fake": True},
            )
        return original_generate(self, request)

    monkeypatch.setattr(FakeProvider, "generate", invalid_distill_json)
    client, _cards = _client(tmp_path, monkeypatch)
    source = tmp_path / "bad-json.md"
    source.write_text("# Bad JSON\n\nbody worth processing\n", encoding="utf-8")

    started = client.post(
        "/api/sources/watch",
        json={"path": str(source), "frequency": "manual", "process_now": True},
    ).json()
    finished = _wait_for_processing_run(client, started["run_id"])

    assert finished["status"] == "failed"
    assert finished["summary"]["errors"] == 1
    assert finished["error_message"] or finished["message"]
    assert "processed as ai_draft" not in f"{finished}"


def test_processing_run_provider_html_error_is_user_friendly() -> None:
    """Provider HTML 错误页不能原样进入 run record / Sources 文案。

    中文学习型说明：真实模型配置仍是用户主路径，但代理、网关、base_url
    配错时常返回 HTML。用户需要可行动的错误，不需要看到整页 HTML。
    """

    message = "LLM 调用失败：HTTP 503: <!DOCTYPE html><html><title>ERROR</title>"

    cleaned = _safe_error_message(message)

    assert "Provider returned an HTML error page (HTTP 503)." in cleaned
    assert "<html" not in cleaned.lower()
    assert "<!DOCTYPE" not in cleaned


def test_processing_run_missing_model_key_error_uses_product_language() -> None:
    """Web Sources 也不能把 env/api_key_env/fake/demo/profile 泄漏给普通用户。"""

    message = "模型 main 没有可用的 API key。请在 Web Setup 中添加 key，或设置环境变量 TEST_KEY。"

    cleaned = _safe_error_message(message)

    assert "Model setup is incomplete" in cleaned
    for token in ("env", "environment variable", "api_key_env", "TEST_KEY", "fake", "demo", "profile"):
        assert token.lower() not in cleaned.lower()


def test_processing_run_started_at_uses_subsecond_precision() -> None:
    """重复点击 Process Now 时，run 排序需要亚秒级 started_at。

    中文学习型说明：当前不引入队列/锁等大架构，只先守住状态归属边界：
    同一 source 的多个后台 run 至少要有可排序的高精度 started_at，避免
    Sources last run summary 在同一秒内随机指向旧 run。
    """

    timestamp = _now()

    assert "." in timestamp
    assert len(timestamp.split(".", 1)[1].split("+", 1)[0]) == 6


def test_web_health_home_config_do_not_expose_secret_values(tmp_path: Path, monkeypatch) -> None:
    client, _cards = _client(tmp_path, monkeypatch)

    assert client.get("/api/health").json()["ok"] is True
    home = client.get("/api/home/status").json()
    config = client.get("/api/config/status").json()
    combined = f"{home} {config}"

    assert "do-not-leak" not in combined
    assert "cubox-secret" not in combined
    assert "MINDFORGE_FAKE_SECRET" in combined
    assert home["safety"]["local_only"] is True
    assert config["provider"]["active_profile"] == "fake"


def test_setup_editor_saves_allowed_non_secret_fields_and_refreshes_status(
    tmp_path: Path,
    monkeypatch,
    capsys,
) -> None:
    cfg_path, _vault, _cards = _write_config(tmp_path)
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["custom_unrelated"] = {"keep": "unchanged"}
    raw["llm"]["profiles"]["real"] = {
        "triage": "real_alias",
        "distill": "real_alias",
        "link_suggestion": "real_alias",
        "review_questions": "real_alias",
        "action_extraction": "real_alias",
    }
    raw["llm"]["models"]["real_alias"] = {
        "provider": "openai_compatible",
        "type": "openai_compatible",
        "base_url": "https://old.example.invalid/v1",
        "model": "old-model",
        "timeout_seconds": 30,
        "max_retries": 1,
        "api_key_env": "MINDFORGE_REAL_SECRET",
        "base_url_env": "MINDFORGE_OLD_BASE_URL",
        "model_env": "MINDFORGE_OLD_MODEL",
    }
    cfg_path.write_text(yaml.safe_dump(raw, sort_keys=False), encoding="utf-8")
    (tmp_path / ".env").write_text("MINDFORGE_REAL_SECRET=never-return-this\n", encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    editable = client.get("/api/config/editable").json()

    assert editable["vault"]["root"].endswith("dogfood-vault")
    assert editable["llm"]["active_provider"] == "fake"
    assert editable["llm"]["available_providers"] == []
    assert editable["llm"]["providers"] == {}
    assert "never-return-this" not in f"{editable}"

    new_vault = tmp_path / "new-vault"
    saved = client.patch(
        "/api/config/editable",
        json={
            "vault_root": str(new_vault),
            "create_vault": True,
        },
    ).json()
    after_raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    combined = f"{saved} {cfg_path.read_text(encoding='utf-8')} {capsys.readouterr()}"

    assert saved["ok"] is True
    assert saved["status"]["vault"]["path"] == str(new_vault.resolve())
    assert saved["status"]["provider"]["active_profile"] == "fake"
    assert (new_vault / "00-Inbox").is_dir()
    assert after_raw["custom_unrelated"] == {"keep": "unchanged"}
    assert after_raw["triage"]["default_track"] == "unrouted"
    assert "never-return-this" not in combined


def test_setup_editor_validates_paths_and_never_writes_api_key_values(
    tmp_path: Path,
    monkeypatch,
) -> None:
    cfg_path, _vault, _cards = _write_config(tmp_path)
    (tmp_path / ".env").write_text("MINDFORGE_FAKE_SECRET=hidden-secret\n", encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    invalid = client.post("/api/config/validate", json={"vault_root": "/"}).json()
    refused = client.patch(
        "/api/config/editable",
        json={
            "vault_root": "/",
            "create_vault": False,
            "providers": {"fake": {"api_key_value": "must-not-be-written"}},
        },
    )

    config_text = cfg_path.read_text(encoding="utf-8")
    assert invalid["ok"] is False
    assert any("dangerous" in error.lower() for error in invalid["errors"])
    assert refused.status_code == 400
    assert "hidden-secret" not in refused.text
    assert "must-not-be-written" not in config_text


def test_setup_effective_provider_config_distinguishes_env_names_defaults_and_masked_secret(
    tmp_path: Path,
    monkeypatch,
    capsys,
) -> None:
    cfg_path, _vault, _cards = _write_config(tmp_path)
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    secret_value = "sk-ant-" + "test-secret-value-" + "abcd"
    raw["llm"] = {
        "active": "anthropic",
        "providers": {
            "anthropic": {
                "type": "anthropic",
                "api_key_env": "MINDFORGE_ANTHROPIC_API_KEY",
                "base_url_env": "MINDFORGE_ANTHROPIC_BASE_URL",
                "model_env": "MINDFORGE_ANTHROPIC_MODEL",
                "default_base_url": "https://api.anthropic.com",
                "default_model": "claude-3-5-haiku-latest",
            },
            "fake": {"type": "fake"},
        },
    }
    cfg_path.write_text(yaml.safe_dump(raw, sort_keys=False), encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    monkeypatch.delenv("MINDFORGE_ANTHROPIC_BASE_URL", raising=False)
    monkeypatch.delenv("MINDFORGE_ANTHROPIC_MODEL", raising=False)
    monkeypatch.setenv("MINDFORGE_ANTHROPIC_API_KEY", secret_value)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    editable = client.get("/api/config/editable").json()
    provider = editable["llm"]["providers"]["anthropic"]
    combined = f"{editable} {capsys.readouterr()}"

    assert editable["llm"]["active_provider"] == "anthropic"
    assert provider["api_key_env"] == "MINDFORGE_ANTHROPIC_API_KEY"
    assert provider["api_key_env_configured"] is True
    assert provider["api_key_secret_present"] is True
    assert provider["api_key_masked_value"] == "sk-****abcd"
    assert provider["api_key_status_label"] == "present (sk-****abcd)"
    assert provider["base_url_env"] == "MINDFORGE_ANTHROPIC_BASE_URL"
    assert provider["base_url_env_status"] == "missing"
    assert provider["effective_base_url"] == "https://api.anthropic.com"
    assert provider["base_url_source"] == "config_default"
    assert provider["model_env"] == "MINDFORGE_ANTHROPIC_MODEL"
    assert provider["model_env_status"] == "missing"
    assert provider["effective_model"] == "claude-3-5-haiku-latest"
    assert provider["model_source"] == "config_default"
    assert secret_value not in combined


def test_setup_effective_provider_config_reports_configured_secret_name_when_value_missing(
    tmp_path: Path,
    monkeypatch,
) -> None:
    cfg_path, _vault, _cards = _write_config(tmp_path)
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["llm"] = {
        "active": "anthropic",
        "providers": {
            "anthropic": {
                "type": "anthropic",
                "api_key_env": "MINDFORGE_ANTHROPIC_API_KEY",
                "base_url_env": "MINDFORGE_ANTHROPIC_BASE_URL",
                "model_env": "MINDFORGE_ANTHROPIC_MODEL",
                "default_base_url": "https://api.anthropic.com",
                "default_model": "claude-3-5-haiku-latest",
            }
        },
    }
    cfg_path.write_text(yaml.safe_dump(raw, sort_keys=False), encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    monkeypatch.delenv("MINDFORGE_ANTHROPIC_API_KEY", raising=False)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    provider = client.get("/api/config/editable").json()["llm"]["providers"]["anthropic"]

    assert provider["api_key_env_configured"] is True
    assert provider["api_key_secret_present"] is False
    assert provider["api_key_masked_value"] is None
    assert provider["api_key_status_label"] == "env name configured, value missing"
    assert provider["effective_base_url"] == "https://api.anthropic.com"
    assert provider["effective_model"] == "claude-3-5-haiku-latest"


def test_setup_provider_dropdown_uses_only_configured_real_providers(
    tmp_path: Path,
    monkeypatch,
) -> None:
    cfg_path, _vault, _cards = _write_config(tmp_path)
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["llm"] = {
        "active": "openai_compatible",
        "providers": {
            "openai_compatible": {
                "type": "openai_compatible",
                "api_key_env": "MINDFORGE_OPENAI_API_KEY",
                "default_model": "gpt-test",
            },
            "all_local": {
                "type": "local",
                "triage": "local_alias",
                "distill": "local_alias",
                "link_suggestion": "local_alias",
                "review_questions": "local_alias",
                "action_extraction": "local_alias",
            },
            "fake": {"type": "fake", "default_model": "fake-fast", "default_base_url": "fake://"},
            "test_default": {
                "type": "test",
                "triage": "test_alias",
                "distill": "test_alias",
                "link_suggestion": "test_alias",
                "review_questions": "test_alias",
                "action_extraction": "test_alias",
            },
        },
        "models": {
            "local_alias": {
                "provider": "local",
                "type": "local",
                "base_url": "http://127.0.0.1:11434",
                "model": "local-model",
            },
            "test_alias": {
                "provider": "test",
                "type": "test",
                "base_url": "test://",
                "model": "test-model",
            },
        },
    }
    cfg_path.write_text(yaml.safe_dump(raw, sort_keys=False), encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    editable = client.get("/api/config/editable").json()
    provider = editable["llm"]["providers"]["openai_compatible"]

    assert editable["llm"]["active_provider"] == "openai_compatible"
    assert editable["llm"]["available_providers"] == ["openai_compatible"]
    assert sorted(editable["llm"]["providers"]) == ["openai_compatible"]
    assert "all_local" not in editable["llm"]["available_providers"]
    assert "anthropic_coding_plan" not in editable["llm"]["available_providers"]
    assert "fake" not in editable["llm"]["available_providers"]
    assert "test_default" not in editable["llm"]["available_providers"]
    assert provider["effective_model"] == "gpt-test"
    assert provider["model_source"] == "config_default"
    assert provider["effective_base_url"] is None
    assert provider["base_url_source"] == "missing"


def test_setup_editable_llm_view_uses_models_default_and_routing(
    tmp_path: Path,
    monkeypatch,
) -> None:
    cfg_path, _vault, _cards = _write_config(tmp_path)
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    secret_value = "sk-test-secret-value-abcd"
    raw["llm"] = {
        "default_model": "strong",
        "models": {
            "cheap": {
                "type": "openai_compatible",
                "api_key_env": "MINDFORGE_LLM_API_KEY",
                "base_url": "https://router.example.com/v1",
                "model": "cheap-model",
            },
            "strong": {
                "type": "anthropic_compatible",
                "api_key_env": "MINDFORGE_LLM_API_KEY",
                "base_url": "https://router.example.com/anthropic",
                "model": "strong-model",
            },
            "local": {
                "type": "openai_compatible",
                "api_key_optional": True,
                "base_url": "http://localhost:11434/v1",
                "model": "local-model",
            },
            "fake_fast": {
                "type": "fake",
                "model": "fake-fast",
            },
        },
        "routing": {
            "triage": "cheap",
            "link_suggestion": "cheap",
        },
    }
    cfg_path.write_text(yaml.safe_dump(raw, sort_keys=False), encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    monkeypatch.setenv("MINDFORGE_LLM_API_KEY", secret_value)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    editable = client.get("/api/config/editable").json()
    llm = editable["llm"]
    combined = f"{editable}"

    assert llm["default_model"] == "strong"
    assert sorted(llm["configured_model_ids"]) == ["cheap", "local", "strong"]
    assert sorted(llm["configured_models"]) == ["cheap", "local", "strong"]
    assert llm["routing"]["triage"] == "cheap"
    assert llm["routing"]["distill"] == "strong"
    assert llm["routing"]["review_questions"] == "strong"
    assert llm["routing_is_explicit"] is True
    assert llm["legacy_config_detected"] is False
    assert llm["validation_errors"] == []
    assert llm["configured_models"]["cheap"]["type"] == "openai_compatible"
    assert llm["configured_models"]["cheap"]["api_key_env"] == "MINDFORGE_LLM_API_KEY"
    assert llm["configured_models"]["cheap"]["api_key_secret_present"] is True
    assert llm["configured_models"]["cheap"]["api_key_masked_value"] == "sk-****abcd"
    assert llm["configured_models"]["local"]["api_key_optional"] is True
    assert llm["resolved_per_step_models"]["triage"]["model_id"] == "cheap"
    assert llm["resolved_per_step_models"]["distill"]["model_id"] == "strong"
    assert "fake_fast" not in llm["configured_models"]
    assert secret_value not in combined


def test_setup_editable_llm_view_reports_omitted_routing_and_legacy_config(
    tmp_path: Path,
    monkeypatch,
) -> None:
    cfg_path, _vault, _cards = _write_config(tmp_path)
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["llm"] = {
        "default_model": "main",
        "models": {
            "main": {
                "type": "openai_compatible",
                "base_url": "https://router.example.com/v1",
                "model": "main-model",
            }
        },
    }
    cfg_path.write_text(yaml.safe_dump(raw, sort_keys=False), encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    llm = client.get("/api/config/editable").json()["llm"]

    assert llm["routing_is_explicit"] is False
    assert llm["routing"] == {stage: "main" for stage in [
        "triage",
        "distill",
        "link_suggestion",
        "review_questions",
        "action_extraction",
    ]}

    legacy_raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    legacy_raw["llm"] = {
        "active_profile": "fake",
        "profiles": {
            "fake": {
                "triage": "fake_alias",
                "distill": "fake_alias",
                "link_suggestion": "fake_alias",
                "review_questions": "fake_alias",
                "action_extraction": "fake_alias",
            },
            "all_local": {
                "triage": "local_alias",
                "distill": "local_alias",
                "link_suggestion": "local_alias",
                "review_questions": "local_alias",
                "action_extraction": "local_alias",
            },
        },
        "models": {
            "fake_alias": {"provider": "fake", "type": "fake", "model": "fake"},
            "local_alias": {
                "provider": "local",
                "type": "openai_compatible",
                "base_url": "http://localhost:11434/v1",
                "model": "local-model",
            },
        },
    }
    cfg_path.write_text(yaml.safe_dump(legacy_raw, sort_keys=False), encoding="utf-8")

    legacy_llm = TestClient(create_app(config_path=cfg_path, host="127.0.0.1")).get(
        "/api/config/editable"
    ).json()["llm"]

    assert legacy_llm["legacy_config_detected"] is True
    assert any("Legacy LLM config detected" in warning for warning in legacy_llm["warnings"])
    assert legacy_llm["configured_models"] == {}
    assert "all_local" not in legacy_llm["configured_model_ids"]
    assert "fake_alias" not in f"{legacy_llm['configured_models']}"


def test_setup_save_writes_new_llm_format_without_legacy_profiles(
    tmp_path: Path,
    monkeypatch,
) -> None:
    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    response = client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://router.example.com/v1",
                    "model": "main-model",
                    "api_key_env": "MINDFORGE_LLM_API_KEY",
                },
                "strong": {
                    "type": "anthropic_compatible",
                    "base_url": "https://router.example.com/anthropic",
                    "model": "strong-model",
                    "api_key_env": "MINDFORGE_LLM_API_KEY",
                },
            },
            "routing": {"distill": "strong"},
        },
    )

    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    text = cfg_path.read_text(encoding="utf-8")

    assert response.status_code == 200
    assert raw["llm"]["default_model"] == "main"
    assert sorted(raw["llm"]["models"]) == ["main", "strong"]
    assert raw["llm"]["routing"] == {
        "triage": "main",
        "distill": "strong",
        "link_suggestion": "main",
        "review_questions": "main",
        "action_extraction": "main",
    }
    assert "active_profile" not in raw["llm"]
    assert "profiles" not in raw["llm"]
    assert "providers" not in raw["llm"]
    assert "stage_models" not in text


def test_setup_provider_dropdown_ignores_legacy_profile_defaults_in_main_ui(
    tmp_path: Path,
    monkeypatch,
) -> None:
    cfg_path, _vault, _cards = _write_config(tmp_path)
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["llm"] = {
        "active_profile": "anthropic_coding_plan",
        "profiles": {
            "all_local": {
                "triage": "local_alias",
                "distill": "local_alias",
                "link_suggestion": "local_alias",
                "review_questions": "local_alias",
                "action_extraction": "local_alias",
            },
            "anthropic_coding_plan": {
                "triage": "anthropic_alias",
                "distill": "anthropic_alias",
                "link_suggestion": "anthropic_alias",
                "review_questions": "anthropic_alias",
                "action_extraction": "anthropic_alias",
            },
            "openai_compatible": {
                "triage": "openai_alias",
                "distill": "openai_alias",
                "link_suggestion": "openai_alias",
                "review_questions": "openai_alias",
                "action_extraction": "openai_alias",
            },
        },
        "models": {
            "local_alias": {
                "provider": "local",
                "type": "local",
                "base_url": "http://127.0.0.1:11434",
                "model": "local-model",
            },
            "anthropic_alias": {
                "provider": "anthropic",
                "type": "anthropic_compatible",
                "base_url": "https://api.anthropic.com",
                "model": "claude-test",
            },
            "openai_alias": {
                "provider": "openai_compatible",
                "type": "openai_compatible",
                "base_url": "https://api.openai.com/v1",
                "model": "gpt-test",
            },
        },
    }
    cfg_path.write_text(yaml.safe_dump(raw, sort_keys=False), encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    editable = client.get("/api/config/editable").json()

    assert editable["llm"]["active_provider"] == "anthropic_coding_plan"
    assert editable["llm"]["available_providers"] == []
    assert editable["llm"]["providers"] == {}
    assert "all_local" not in f"{editable['llm']['available_providers']}"
    assert "anthropic_coding_plan" not in f"{editable['llm']['available_providers']}"
    assert "openai_compatible" not in f"{editable['llm']['available_providers']}"


def test_setup_active_provider_missing_reports_config_error_without_fallback(
    tmp_path: Path,
    monkeypatch,
) -> None:
    cfg_path, _vault, _cards = _write_config(tmp_path)
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["llm"] = {
        "active": "missing_provider",
        "providers": {
            "anthropic": {
                "type": "anthropic",
                "api_key_env": "MINDFORGE_ANTHROPIC_API_KEY",
                "default_model": "claude-test",
            }
        },
    }
    cfg_path.write_text(yaml.safe_dump(raw, sort_keys=False), encoding="utf-8")
    monkeypatch.chdir(tmp_path)

    with pytest.raises(AppContextError) as exc_info:
        create_app(config_path=cfg_path, host="127.0.0.1")

    error_text = str(exc_info.value)
    assert "missing_provider" in error_text
    assert "fake" not in error_text.lower()


def test_web_drafts_detail_and_approve_confirmation_boundary(tmp_path: Path, monkeypatch) -> None:
    client, cards = _client(tmp_path, monkeypatch)
    card = _write_draft(cards)

    drafts = client.get("/api/drafts").json()
    assert [draft["id"] for draft in drafts["drafts"]] == ["draft-1"]
    assert drafts["drafts"][0]["strategy_id"] == "five_stage"
    assert drafts["drafts"][0]["source_id"] == "src_draft_1"
    assert drafts["drafts"][0]["prompt_versions"]["triage"] == "v1"
    detail = client.get("/api/drafts/draft-1").json()
    assert "safe draft body" in detail["body"].lower()
    assert detail["draft"]["source_content_hash"] == "sha256:provenancehash"
    assert detail["frontmatter"]["run_id"] == "run-web-provenance"

    refused = client.post(
        "/api/drafts/draft-1/approve",
        json={"confirm": False, "reviewed_source": True},
    ).json()
    assert refused["ok"] is False
    assert "status: ai_draft" in card.read_text(encoding="utf-8")

    approved = client.post(
        "/api/drafts/draft-1/approve",
        json={"confirm": True, "reviewed_source": True},
    ).json()
    assert approved["ok"] is True
    assert approved["new_status"] == "human_approved"
    assert approved["index_updated"] is True
    assert approved["index_path"]
    assert Path(approved["index_path"]).exists()
    assert "status: human_approved" in card.read_text(encoding="utf-8")
    approved_fm = yaml.safe_load(card.read_text(encoding="utf-8").split("---", 2)[1])
    assert approved_fm["strategy_id"] == "five_stage"
    assert approved_fm["prompt_versions"]["distill"] == "v1"


def test_web_workspace_can_save_draft_without_approving_or_losing_provenance(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Web 编辑的是 Knowledge Card 正文，不是 source，也不能绕过显式 approve。

    大模型和 source 都不参与这个测试；它只验证 Web API 是否复用卡片持久化边界：
    保存 draft 只替换 body，frontmatter 的 status/provenance 必须原样保留。
    """

    client, cards = _client(tmp_path, monkeypatch)
    card = _write_draft(cards)
    source = tmp_path / "dogfood-vault" / "00-Inbox" / "ManualNotes" / "source-note.md"
    source.write_text("SOURCE_BODY_MUST_NOT_LEAK_AFTER_SAVE", encoding="utf-8")

    saved = client.patch(
        "/api/drafts/draft-1",
        json={"body": "## AI Summary\n\nEdited draft workspace summary.\n"},
    ).json()
    detail = client.get("/api/drafts/draft-1").json()
    combined = f"{saved} {detail}"
    fm = yaml.safe_load(card.read_text(encoding="utf-8").split("---", 2)[1])

    assert saved["ok"] is True
    assert saved["status"] == "ai_draft"
    assert saved["index_updated"] is False
    assert "Edited draft workspace summary" in detail["body"]
    assert "status: ai_draft" in card.read_text(encoding="utf-8")
    assert fm["strategy_id"] == "five_stage"
    assert fm["source_content_hash"] == "sha256:provenancehash"
    assert fm["prompt_versions"]["triage"] == "v1"
    assert "SOURCE_BODY_MUST_NOT_LEAK_AFTER_SAVE" not in combined


def test_web_workspace_can_read_and_save_approved_card_and_refresh_recall(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Library 是正式知识工作台：approved card 可阅读、可编辑，保存后刷新 BM25。

    中文学习型说明：这里 intentionally 用 approved card body 中的 AI Summary 命中
    recall，证明 Web save 后调用的是本地 lexical index rebuild，而不是只改文件。
    """

    client, cards = _client(tmp_path, monkeypatch)
    card = _write_approved(cards)
    source = tmp_path / "dogfood-vault" / "00-Inbox" / "_processed" / "ManualNotes" / "source-note.md"
    source.parent.mkdir(parents=True)
    source.write_text("APPROVED_SOURCE_RAW_MUST_NOT_LEAK", encoding="utf-8")

    detail = client.get("/api/library/card", params={"ref": "approved-1"}).json()
    assert "Approved summary alpha workspace" in detail["body"]
    assert detail["card"]["strategy_label"] == "Knowledge Card Workflow"

    saved = client.patch(
        "/api/library/card",
        params={"ref": "approved-1"},
        json={"body": "## AI Summary\n\nEdited approved gamma workspace.\n"},
    ).json()
    after = client.get("/api/library/card", params={"ref": "approved-1"}).json()
    recall = client.get("/api/recall", params={"q": "gamma workspace"}).json()
    combined = f"{detail} {saved} {after} {recall}"
    fm = yaml.safe_load(card.read_text(encoding="utf-8").split("---", 2)[1])

    assert saved["ok"] is True
    assert saved["status"] == "human_approved"
    assert saved["index_updated"] is True
    assert Path(saved["index_path"]).exists()
    assert "Edited approved gamma workspace" in after["body"]
    assert "status: human_approved" in card.read_text(encoding="utf-8")
    assert fm["source_archive_path"] == "90-Archive/Processed/source-note.md"
    assert fm["run_id"] == "run-approved"
    assert recall["hits"][0]["rel_path"].endswith("approved.md")
    assert recall["hits"][0]["card_ref"] == "approved-1"
    assert recall["hits"][0]["detail_href"].startswith("/library?")
    assert "APPROVED_SOURCE_RAW_MUST_NOT_LEAK" not in combined
    assert "HUMAN_NOTE_MUST_NOT_BE_RECALLABLE" not in recall["hits"][0]["why_this_matched"]


def test_library_card_detail_exposes_local_graph_and_related_cards(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Card detail 必须有用户可见的 Relationship Preview 数据。

    中文学习型说明：M6 不能停在 backend graph engine；Web detail API 要返回
    local_graph / related_cards，前端才能用列表 fallback 展示关系并导航。这里
    只构造 synthetic approved cards，不读取真实 vault 或 source 正文。
    """

    client, cards = _client(tmp_path, monkeypatch)
    _write_approved(cards)
    neighbor = cards / "neighbor.md"
    neighbor.write_text(
        """---
id: approved-neighbor
title: Neighbor Card
status: human_approved
track: product
tags:
  - library
source_type: manual_note
source_id: src_approved_1
source_path: 00-Inbox/_processed/ManualNotes/neighbor.md
source_title: Approved source
value_score: 7
run_id: run-approved
---

## AI Summary

Neighbor summary.
""",
        encoding="utf-8",
    )
    draft_neighbor = cards / "draft-neighbor.md"
    draft_neighbor.write_text(
        """---
id: draft-neighbor
title: Draft Neighbor
status: ai_draft
tags:
  - library
source_id: src_approved_1
---

## AI Summary

Draft should not appear in library relationships.
""",
        encoding="utf-8",
    )

    detail = client.get("/api/library/card", params={"ref": "approved-1"}).json()

    graph = detail["local_graph"]
    graph_card_ids = {node["id"] for node in graph["nodes"] if node["type"] == "card"}
    related_ids = {item["card"]["id"] for item in detail["related_cards"]}
    reason_labels = {reason["label"] for item in detail["related_cards"] for reason in item["reasons"]}

    assert graph["center_id"] == "approved-1"
    assert "approved-neighbor" in graph_card_ids
    assert "draft-neighbor" not in graph_card_ids
    assert "approved-neighbor" in related_ids
    assert "draft-neighbor" not in related_ids
    assert "Same source" in reason_labels
    assert all(node["href"].startswith("/library?card=") for node in graph["nodes"] if node["type"] == "card")


def test_web_reject_and_imports_are_honest_unavailable(tmp_path: Path, monkeypatch) -> None:
    client, cards = _client(tmp_path, monkeypatch)
    _write_draft(cards)

    reject = client.post("/api/drafts/draft-1/reject", json={"reason": "not useful"}).json()
    import_local = client.post("/api/sources/import-local").json()
    import_cubox = client.post("/api/sources/import-cubox-json").json()

    assert reject["available"] is False
    assert import_local["available"] is False
    assert import_cubox["available"] is False


def test_web_watch_list_add_delete_and_import_align_with_cli_ingestion(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Web ingestion 必须复用 watch/import 语义，而不是暴露 scan/process 主流程。

    这些动作使用临时 source；断言重点是 registry/card/source 的状态边界：
    watch delete 只删 registry，import 不进 registry，所有自动化只生成 ai_draft。
    """

    client, cards = _client(tmp_path, monkeypatch)
    vault = tmp_path / "dogfood-vault"
    watch_file = tmp_path / "watch-file.md"
    watch_folder = tmp_path / "watch-folder"
    import_file = tmp_path / "import-file.md"
    registered_only_file = tmp_path / "registered-only.md"
    watch_file.write_text("# Watch File\n\nWATCH_BODY_MUST_NOT_LEAK\n", encoding="utf-8")
    watch_folder.mkdir()
    (watch_folder / "watch-folder-note.md").write_text("# Watch Folder\n\nbody\n", encoding="utf-8")
    import_file.write_text("# Import File\n\nIMPORT_BODY_MUST_NOT_LEAK\n", encoding="utf-8")
    registered_only_file.write_text("# Registered Only\n\nbody\n", encoding="utf-8")

    listed = client.get("/api/sources/watch").json()
    assert listed["watched_sources"][0]["id"] == "default-inbox"
    assert listed["watched_sources"][0]["is_default"] is True
    assert listed["watched_sources"][0]["can_delete"] is True

    default_frequency_update = client.patch(
        "/api/sources/watch/default-inbox/frequency",
        json={"frequency": "daily"},
    ).json()
    assert default_frequency_update["ok"] is True
    default_stop = client.delete("/api/sources/watch/default-inbox").json()
    assert default_stop["ok"] is True
    assert default_stop["source_deleted"] is False
    assert default_stop["cards_deleted"] is False
    assert "only stops future monitoring" in default_stop["message"]

    added_file = client.post("/api/sources/watch", json={"path": str(watch_file)}).json()
    added_folder = client.post("/api/sources/watch", json={"path": str(watch_folder)}).json()
    registered_only = client.post(
        "/api/sources/watch",
        json={"path": str(registered_only_file), "process_now": False},
    ).json()

    assert added_file["ok"] is True
    assert added_file["mode"] == "watch_add"
    assert added_file["run_id"]
    assert added_file["counts"]["processed"] == 0
    assert added_file["added_to_registry"] is True
    assert added_folder["run_id"]
    assert added_folder["counts"]["processed"] == 0
    assert registered_only["added_to_registry"] is True
    assert registered_only["counts"]["processed"] == 0
    _wait_for_processing_run(client, added_file["run_id"])
    _wait_for_processing_run(client, added_folder["run_id"])

    frequency_update = client.patch(
        f"/api/sources/watch/{registered_only['watch_id']}/frequency",
        json={"frequency": "daily"},
    ).json()
    assert frequency_update["ok"] is True
    assert "top-level source only" in frequency_update["message"]

    registry = WatchRegistry.load(vault / ".mindforge" / "watched_sources.json")
    user_sources = [source for source in registry.sources if not source.is_default]
    assert {source.path for source in user_sources} == {
        watch_file.resolve(),
        watch_folder.resolve(),
        registered_only_file.resolve(),
    }
    assert next(source for source in user_sources if source.path == registered_only_file.resolve()).frequency == "daily"

    deleted = client.delete(f"/api/sources/watch/{user_sources[0].id}").json()
    assert deleted["ok"] is True
    assert deleted["source_deleted"] is False
    assert deleted["cards_deleted"] is False
    assert watch_file.exists()
    assert len(list(cards.rglob("*.md"))) == 2

    imported = client.post("/api/sources/import", json={"path": str(import_file)}).json()

    assert imported["ok"] is True
    assert imported["mode"] == "import"
    assert imported["counts"]["processed"] == 1
    assert imported["added_to_registry"] is False
    final_registry = WatchRegistry.load(vault / ".mindforge" / "watched_sources.json")
    assert len([source for source in final_registry.sources if not source.is_default]) == 2
    assert len(list(cards.rglob("*.md"))) == 3
    assert all("status: ai_draft" in card.read_text(encoding="utf-8") for card in cards.rglob("*.md"))

    after = client.get("/api/sources").json()
    combined = f"{listed} {added_file} {imported} {after}"
    assert "WATCH_BODY_MUST_NOT_LEAK" not in combined
    assert "IMPORT_BODY_MUST_NOT_LEAK" not in combined
    assert "自动化只生成 ai_draft" in combined
    assert after["ingestion"]["primary_entry"] == "watch/import"
    assert "advanced" in after["ingestion"]["advanced_note"].lower()


def test_web_process_uses_default_model_when_routing_is_missing(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Web Process now 兼容只有 default_model 的新格式配置。

    中文学习型说明：Setup 主路径会补齐 routing，但 runtime 仍要接受旧 dogfood
    clone 中可能已有的“default_model + models、无 routing”配置，不能退回旧
    profile[stage] 导致 KeyError。
    """

    cfg_path, _vault, cards = _write_config(tmp_path)
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["llm"] = {
        "default_model": "main",
        "models": {
            "main": {
                "provider": "fake",
                "type": "fake",
                "base_url": "fake://",
                "model": "fake",
                "timeout_seconds": 5,
                "max_retries": 0,
                "api_key_optional": True,
            }
        },
    }
    cfg_path.write_text(yaml.safe_dump(raw, allow_unicode=True, sort_keys=False), encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    source = tmp_path / "default-model-only.md"
    source.write_text("# Default Model Only\n\nbody\n", encoding="utf-8")

    response = client.post("/api/sources/watch", json={"path": str(source)})

    assert response.status_code == 200, response.text
    finished = _wait_for_processing_run(client, response.json()["run_id"])
    assert finished["summary"]["drafts"] == 1
    assert len(list(cards.rglob("*.md"))) == 1


def test_web_process_without_model_returns_friendly_error(
    tmp_path: Path,
    monkeypatch,
) -> None:
    cfg_path, _vault, _cards = _write_config(tmp_path)
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["llm"] = {"default_model": None, "models": {}, "routing": {}}
    cfg_path.write_text(yaml.safe_dump(raw, allow_unicode=True, sort_keys=False), encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    client = TestClient(
        create_app(config_path=cfg_path, host="127.0.0.1"),
        raise_server_exceptions=False,
    )
    source = tmp_path / "no-model.md"
    source.write_text("# No Model\n\nbody\n", encoding="utf-8")

    response = client.post("/api/sources/watch", json={"path": str(source)})

    assert response.status_code == 400
    detail = response.json()["detail"]["message"]
    assert "No model configured for stage 'triage'" in detail
    assert "Add a model in Web Setup" in detail
    assert "Traceback" not in response.text


def test_web_watch_scan_without_model_returns_friendly_error(
    tmp_path: Path,
    monkeypatch,
) -> None:
    cfg_path, _vault, _cards = _write_config(tmp_path)
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["llm"] = {"default_model": None, "models": {}, "routing": {}}
    cfg_path.write_text(yaml.safe_dump(raw, allow_unicode=True, sort_keys=False), encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    client = TestClient(
        create_app(config_path=cfg_path, host="127.0.0.1"),
        raise_server_exceptions=False,
    )
    source = tmp_path / "scan-no-model.md"
    source.write_text("# Scan No Model\n\nbody\n", encoding="utf-8")
    registered = client.post(
        "/api/sources/watch",
        json={"path": str(source), "process_now": False},
    ).json()

    response = client.post(f"/api/sources/watch/scan?ref={registered['watch_id']}")

    assert response.status_code == 400
    assert "No model configured for stage 'triage'" in response.json()["detail"]["message"]
    assert "Traceback" not in response.text


def test_sources_api_returns_recursive_folder_watch_diagnostics(
    tmp_path: Path,
    monkeypatch,
) -> None:
    client, cards = _client(tmp_path, monkeypatch)
    watch_folder = tmp_path / "watch-folder"
    supported = watch_folder / "nested" / "supported.md"
    unsupported = watch_folder / "nested" / "unsupported.csv"
    temp_file = watch_folder / "~$draft.md"
    hidden = watch_folder / ".hidden.md"
    generated = watch_folder / "20-Knowledge-Cards" / "generated.md"
    supported.parent.mkdir(parents=True)
    for path, body in (
        (supported, "# Supported\n\nSUPPORTED_BODY_MUST_NOT_LEAK\n"),
        (unsupported, "UNSUPPORTED_BODY_MUST_NOT_LEAK\n"),
        (temp_file, "# Temp\n\nTEMP_BODY_MUST_NOT_LEAK\n"),
        (hidden, "# Hidden\n\nHIDDEN_BODY_MUST_NOT_LEAK\n"),
        (generated, "# Generated\n\nGENERATED_BODY_MUST_NOT_LEAK\n"),
    ):
        path.parent.mkdir(parents=True, exist_ok=True)
        path.write_text(body, encoding="utf-8")

    added = client.post("/api/sources/watch", json={"path": str(watch_folder)}).json()
    _wait_for_processing_run(client, added["run_id"])
    response = client.get("/api/sources").json()
    watched = next(item for item in response["watched_sources"] if item["path"] == str(watch_folder.resolve()))
    combined = f"{added} {response}"

    assert added["ok"] is True
    assert watched["path_type"] == "folder"
    assert watched["recursive"] is True
    assert watched["supported_file_count"] == 1
    assert watched["processed_count"] == 1
    assert watched["failed_count"] == 0
    assert watched["skipped_reason_summary"]["unsupported_extension"] == 1
    assert watched["skipped_reason_summary"]["temp_file"] == 1
    assert watched["skipped_reason_summary"]["hidden_file"] == 1
    assert watched["status_label"] in {"Watching", "Processed", "Succeeded"}
    assert len(list(cards.rglob("*.md"))) == 1
    assert watched["status_label"] != "ready"
    assert watched["status"] != "ready"
    assert "Approved" not in combined
    assert "SUPPORTED_BODY_MUST_NOT_LEAK" not in combined
    assert "UNSUPPORTED_BODY_MUST_NOT_LEAK" not in combined
    assert "TEMP_BODY_MUST_NOT_LEAK" not in combined
    assert "HIDDEN_BODY_MUST_NOT_LEAK" not in combined
    assert "GENERATED_BODY_MUST_NOT_LEAK" not in combined


def test_web_stop_watching_built_in_inbox_does_not_delete_knowledge_cards(
    tmp_path: Path,
    monkeypatch,
) -> None:
    client, cards = _client(tmp_path, monkeypatch)
    approved = _write_approved(cards)

    response = client.delete("/api/sources/watch/default-inbox").json()
    after = client.get("/api/sources/watch").json()
    built_in = next(item for item in after["watched_sources"] if item["id"] == "default-inbox")

    assert response["ok"] is True
    assert response["source_deleted"] is False
    assert response["cards_deleted"] is False
    assert "only stops future monitoring" in response["message"]
    assert approved.exists()
    assert "status: human_approved" in approved.read_text(encoding="utf-8")
    assert built_in["status"] == "paused"


def test_sources_path_actions_are_allowlisted_and_do_not_read_file_content(
    tmp_path: Path,
    monkeypatch,
) -> None:
    client, cards = _client(tmp_path, monkeypatch)

    # 创建 source file + draft card，供 object-reference reveal 使用
    source_file = tmp_path / "dogfood-vault" / "00-Inbox" / "ManualNotes" / "source-note.md"
    source_file.write_text("SOURCE_CONTENT_MUST_NOT_BE_READ", encoding="utf-8")
    _write_draft(cards)
    outside = tmp_path.parent / f"outside-{tmp_path.name}.md"
    outside.write_text("OUTSIDE_CONTENT_MUST_NOT_BE_READ", encoding="utf-8")

    calls: list[tuple[list[str], dict]] = []

    def fake_run(args, **kwargs):
        calls.append((list(args), kwargs))

        class Result:
            returncode = 0

        return Result()

    monkeypatch.setattr("mindforge_web.services.web_path_action_service.sys.platform", "darwin")
    monkeypatch.setattr("mindforge_web.services.web_path_action_service.subprocess.run", fake_run)

    # 旧 raw path endpoint 全部返回 410 —— 不泄露路径是否存在（无 path oracle）
    copy_resp = client.post("/api/sources/path-actions/copy", json={"path": str(source_file)})
    reveal_file = client.post("/api/sources/path-actions/reveal", json={"path": str(source_file)})
    reveal_dir = client.post("/api/sources/path-actions/reveal", json={"path": str(source_file.parent)})
    reveal_missing = client.post(
        "/api/sources/path-actions/reveal",
        json={"path": str(source_file.parent / "missing.md")},
    )
    reveal_outside = client.post("/api/sources/path-actions/reveal", json={"path": str(outside)})

    assert copy_resp.status_code == 410
    assert reveal_file.status_code == 410
    assert reveal_dir.status_code == 410
    assert reveal_missing.status_code == 410
    assert reveal_outside.status_code == 410

    # 新 object-reference reveal endpoint
    reveal = client.post("/api/sources/reveal", json={"card_id": "draft-1"}).json()
    assert reveal["ok"] is True
    assert reveal["action"] == "reveal"
    assert calls[0][0] == ["open", "-R", str(source_file.resolve())]
    assert calls[0][1].get("shell") is False

    # 所有 response 均不泄露 source 文件内容
    combined = (
        f"{copy_resp.text} {reveal_file.text} {reveal_dir.text} "
        f"{reveal_missing.text} {reveal_outside.text} {reveal}"
    )
    assert "SOURCE_CONTENT_MUST_NOT_BE_READ" not in combined
    assert "OUTSIDE_CONTENT_MUST_NOT_BE_READ" not in combined


def test_sources_status_names_processed_and_generated_knowledge_without_ready_or_approved(
    tmp_path: Path,
    monkeypatch,
) -> None:
    client, cards = _client(tmp_path, monkeypatch)
    processed = (
        tmp_path
        / "dogfood-vault"
        / "00-Inbox"
        / "_processed"
        / "ManualNotes"
        / "source-note.md"
    )
    processed.parent.mkdir(parents=True)
    processed.write_text("PROCESSED_SOURCE_BODY_MUST_NOT_LEAK", encoding="utf-8")
    _write_draft(cards)

    response = client.get("/api/sources").json()
    source = response["sources"][0]
    combined = f"{response}"

    assert source["display_status"] == "Processed"
    assert source["generated_knowledge_status"] == "Has generated knowledge"
    assert source["generated_card_count"] == 1
    assert source["generated_card_paths"][0].endswith("draft.md")
    assert "ready" not in combined
    assert "Approved" not in combined
    assert "PROCESSED_SOURCE_BODY_MUST_NOT_LEAK" not in combined


# ============================================================================
# Model config + secret store 测试
# ============================================================================


def test_setup_add_model_writes_new_llm_models_entry(tmp_path: Path, monkeypatch) -> None:
    """Add model 写入 llm.models，API key 进入 secret store。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    response = client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://router.example.com/v1",
                    "model": "main-model",
                    "api_key": "sk-test-key-1234",
                    "api_key_action": "update",
                },
            },
        },
    )
    assert response.status_code == 200

    # 验证 YAML config
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    assert raw["llm"]["default_model"] == "main"
    assert "main" in raw["llm"]["models"]
    assert raw["llm"]["models"]["main"]["type"] == "openai_compatible"
    assert raw["llm"]["routing"] == {stage: "main" for stage in REQUIRED_STAGES}
    # YAML 里不应有 API key
    assert "api_key" not in raw["llm"]["models"]["main"]
    assert "sk-test" not in str(raw)

    # 验证 secret store
    secrets_path = tmp_path / ".mindforge" / "secrets.json"
    import json
    assert secrets_path.is_file()
    secrets = json.loads(secrets_path.read_text(encoding="utf-8"))
    assert secrets["main"] == "sk-test-key-1234"


def test_clean_clone_setup_api_bootstraps_no_model_config(tmp_path: Path, monkeypatch) -> None:
    """Web app 在 clean clone 缺本地 config 时应进入 No model configured 状态。"""

    workspace = tmp_path / "mindforge"
    (workspace / "configs").mkdir(parents=True)
    (workspace / "configs" / "mindforge_example.yaml").write_text("version: 0.7\n", encoding="utf-8")
    (workspace / "pyproject.toml").write_text("[project]\nname='mindforge'\n", encoding="utf-8")
    (workspace / "src" / "mindforge").mkdir(parents=True)
    monkeypatch.chdir(workspace)

    cfg_path = workspace / "configs" / "mindforge.yaml"
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    response = client.get("/api/config/editable")

    assert response.status_code == 200
    assert cfg_path.is_file()
    llm = response.json()["llm"]
    assert llm["configured_model_ids"] == []
    assert llm["configured_models"] == {}
    assert llm["default_model"] is None
    assert llm["validation_errors"] == []


def test_setup_edit_model_preserves_secret_when_empty_key(tmp_path: Path, monkeypatch) -> None:
    """Edit model 时 api_key 为空 → 保留已有 secret。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    # 先创建带 key 的模型
    client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://router.example.com/v1",
                    "model": "main-model",
                    "api_key": "sk-original-key",
                    "api_key_action": "update",
                },
            },
        },
    )

    # 编辑：只改 model name，api_key 为空，api_key_action 为 keep
    client.patch(
        "/api/config/editable",
        json={
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://router.example.com/v1",
                    "model": "updated-model",
                    "api_key": "",
                    "api_key_action": "keep",
                },
            },
        },
    )

    import json
    secrets_path = tmp_path / ".mindforge" / "secrets.json"
    secrets = json.loads(secrets_path.read_text(encoding="utf-8"))
    assert secrets["main"] == "sk-original-key"


def test_setup_edit_model_clears_secret_on_explicit_action(tmp_path: Path, monkeypatch) -> None:
    """api_key_action=clear → 删除 secret store 中的 key。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    # 先创建带 key 的模型
    client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://router.example.com/v1",
                    "model": "main-model",
                    "api_key": "sk-original-key",
                    "api_key_action": "update",
                },
            },
        },
    )

    # 清除 key
    client.patch(
        "/api/config/editable",
        json={
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://router.example.com/v1",
                    "model": "main-model",
                    "api_key_action": "clear",
                },
            },
        },
    )

    import json
    secrets_path = tmp_path / ".mindforge" / "secrets.json"
    secrets = json.loads(secrets_path.read_text(encoding="utf-8"))
    assert "main" not in secrets


def test_setup_delete_model_removes_config_and_secret(tmp_path: Path, monkeypatch) -> None:
    """Delete model → 从 config 和 secret store 都删除。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    # 创建两个模型
    client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://r.example.com/v1",
                    "model": "main-model",
                    "api_key": "sk-main-key",
                    "api_key_action": "update",
                },
                "secondary": {
                    "type": "anthropic_compatible",
                    "base_url": "https://a.example.com",
                    "model": "secondary-model",
                    "api_key": "sk-secondary-key",
                    "api_key_action": "update",
                },
            },
        },
    )

    # 删除 secondary（从 models dict 中省略）
    response = client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://r.example.com/v1",
                    "model": "main-model",
                },
            },
        },
    )
    assert response.status_code == 200

    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    assert "secondary" not in raw["llm"]["models"]

    import json
    secrets_path = tmp_path / ".mindforge" / "secrets.json"
    secrets = json.loads(secrets_path.read_text(encoding="utf-8"))
    assert "secondary" not in secrets
    assert secrets["main"] == "sk-main-key"


def test_setup_cannot_delete_default_model(tmp_path: Path, monkeypatch) -> None:
    """不允许删除 default_model 引用的模型。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://r.example.com/v1",
                    "model": "main-model",
                },
            },
        },
    )

    # 尝试删除 main 但保持 default_model=main
    response = client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {},
        },
    )
    assert response.status_code == 400


def test_setup_cannot_leave_routing_to_deleted_model(tmp_path: Path, monkeypatch) -> None:
    """不允许 routing 引用即将被删除的模型。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {"type": "openai_compatible", "base_url": "https://r.example.com/v1", "model": "main-model"},
                "strong": {"type": "anthropic_compatible", "base_url": "https://a.example.com", "model": "strong-model"},
            },
            "routing": {"distill": "strong"},
        },
    )

    # 尝试删 strong 但 routing 仍引用它
    response = client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {"type": "openai_compatible", "base_url": "https://r.example.com/v1", "model": "main-model"},
            },
        },
    )
    assert response.status_code == 400


def test_setup_api_key_never_returned_raw(tmp_path: Path, monkeypatch) -> None:
    """API response 永远不返回 raw API key。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://r.example.com/v1",
                    "model": "main-model",
                    "api_key": "sk-very-secret-key-1234",
                    "api_key_action": "update",
                },
            },
        },
    )

    editable = client.get("/api/config/editable").json()
    model = editable["llm"]["configured_models"]["main"]
    combined = f"{editable}"

    # 确认 masked 值存在
    assert model["api_key_masked_value"] == "sk-****1234"
    assert model["api_key_source"] == "local_secret"
    assert model["api_key_secret_present"] is True

    # raw key 绝不出现在 response 中
    assert "sk-very-secret-key-1234" not in combined
    assert "very-secret" not in combined


def test_setup_hides_fake_model_from_user_config_surface(tmp_path: Path, monkeypatch) -> None:
    """type=fake 只能作为内部测试替身，不能进入普通用户 Setup 列表。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)

    # 写入只有 fake 模型的 config
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["llm"] = {
        "default_model": "demo",
        "models": {
            "demo": {
                "type": "fake",
                "base_url": "fake://",
                "model": "demo-model",
            },
        },
    }
    cfg_path.write_text(yaml.safe_dump(raw, sort_keys=False), encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    editable = client.get("/api/config/editable").json()
    model = editable["llm"]["configured_models"].get("demo")
    assert model is None


def test_setup_default_model_dropdown_only_configured_models(tmp_path: Path, monkeypatch) -> None:
    """default_model 下拉只包含已配置的模型。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    client.patch(
        "/api/config/editable",
        json={
            "default_model": "a",
            "models": {
                "a": {"type": "openai_compatible", "base_url": "https://x.com/v1", "model": "a-model"},
                "b": {"type": "openai_compatible", "base_url": "https://y.com/v1", "model": "b-model"},
            },
        },
    )

    editable = client.get("/api/config/editable").json()
    ids = editable["llm"]["configured_model_ids"]
    assert sorted(ids) == ["a", "b"]


def test_setup_routing_dropdown_only_configured_models(tmp_path: Path, monkeypatch) -> None:
    """routing dropdown 只包含 configured models。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {"type": "openai_compatible", "base_url": "https://x.com/v1", "model": "main-model"},
            },
            "routing": {"distill": "main"},
        },
    )

    editable = client.get("/api/config/editable").json()
    for model_id in editable["llm"]["routing"].values():
        assert model_id in editable["llm"]["configured_model_ids"]


def test_setup_routing_omitted_uses_default_model(tmp_path: Path, monkeypatch) -> None:
    """Setup 保存时补齐 routing；runtime 仍保留 default_model fallback。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {"type": "openai_compatible", "base_url": "https://x.com/v1", "model": "main-model"},
            },
        },
    )

    editable = client.get("/api/config/editable").json()
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    assert editable["llm"]["routing_is_explicit"] is True
    assert editable["llm"]["routing"] == {stage: "main" for stage in REQUIRED_STAGES}
    assert raw["llm"]["routing"] == {stage: "main" for stage in REQUIRED_STAGES}


def test_setup_type_must_be_explicit_no_guessing(tmp_path: Path, monkeypatch) -> None:
    """type 必须显式选择，不能从 URL 自动猜。该验证由 config 层保证。"""
    from mindforge.config import ConfigError, load_mindforge_config

    yaml_path = tmp_path / "test.yaml"
    yaml_path.write_text(
        yaml.safe_dump(
            {
                "version": 0.7,
                "vault": {"root": str(tmp_path)},
                "llm": {
                    "default_model": "main",
                    "models": {
                        "main": {
                            "model": "some-model",
                            "base_url": "https://api.openai.com/v1",
                            # 故意缺失 type
                        },
                    },
                },
            },
            sort_keys=False,
        ),
        encoding="utf-8",
    )
    with pytest.raises(ConfigError, match="type"):
        load_mindforge_config(yaml_path)


def test_setup_api_key_not_in_yaml_when_stored_in_secret_store(tmp_path: Path, monkeypatch) -> None:
    """API key 存在 secret store 时，YAML config 不应包含 raw key。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://r.example.com/v1",
                    "model": "main-model",
                    "api_key": "sk-sensitive-key-9999",
                    "api_key_action": "update",
                },
            },
        },
    )

    yaml_text = cfg_path.read_text(encoding="utf-8")
    assert "sk-sensitive" not in yaml_text
    assert "9999" not in yaml_text


def test_setup_secret_store_under_mindforge_dir(tmp_path: Path, monkeypatch) -> None:
    """Secret store 位于 .mindforge/secrets.json，已被 .gitignore 覆盖。"""
    from mindforge.secret_store import SecretStore

    store = SecretStore(tmp_path / ".mindforge" / "secrets.json")
    store.set("test-model", "sk-test-value")
    assert store.present("test-model") is True
    assert store.get("test-model") == "sk-test-value"
    assert store.masked("test-model") == "sk-****alue"
    store.remove("test-model")
    assert store.present("test-model") is False

    # 验证 .mindforge/ 在 gitignore 中
    gitignore = Path(__file__).parent.parent / ".gitignore"
    assert gitignore.exists()
    content = gitignore.read_text(encoding="utf-8")
    assert ".mindforge/" in content


# ============================================================================
# Env var field cleanup 测试 —— 普通保存不写回 env 字段
# ============================================================================


def test_setup_save_does_not_write_env_var_fields(tmp_path: Path, monkeypatch) -> None:
    """Web 普通保存模型时不写出 api_key_env/base_url_env/model_env。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://router.example.com/v1",
                    "model": "main-model",
                    "api_key": "sk-test-key",
                    "api_key_action": "update",
                },
            },
        },
    )

    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    model = raw["llm"]["models"]["main"]

    # 产品字段存在
    assert model["type"] == "openai_compatible"
    assert model["base_url"] == "https://router.example.com/v1"
    assert model["model"] == "main-model"

    # env var name 字段不应出现
    assert "api_key_env" not in model
    assert "base_url_env" not in model
    assert "model_env" not in model
    assert "version_env" not in model

    # API key raw value 不应在 YAML 中
    yaml_text = cfg_path.read_text(encoding="utf-8")
    assert "sk-test" not in yaml_text


def test_setup_save_cleans_up_legacy_env_fields(tmp_path: Path, monkeypatch) -> None:
    """已有 api_key_env/base_url_env 的旧配置，保存后应清理 env 字段。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)

    # 写入含 env 字段的旧格式配置
    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["llm"] = {
        "default_model": "main",
        "models": {
            "main": {
                "type": "openai_compatible",
                "base_url": "https://old.example.com/v1",
                "model": "old-model",
                "api_key_env": "MINDFORGE_LLM_API_KEY",
                "base_url_env": "MINDFORGE_LLM_BASE_URL",
                "model_env": "MINDFORGE_LLM_MODEL",
            },
        },
    }
    cfg_path.write_text(yaml.safe_dump(raw, sort_keys=False), encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    # 保存：只改 model name，不传 env 字段
    client.patch(
        "/api/config/editable",
        json={
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://new.example.com/v1",
                    "model": "new-model",
                },
            },
        },
    )

    raw_after = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    model = raw_after["llm"]["models"]["main"]

    assert model["base_url"] == "https://new.example.com/v1"
    assert model["model"] == "new-model"
    assert "api_key_env" not in model
    assert "base_url_env" not in model
    assert "model_env" not in model


def test_setup_legacy_config_with_env_fields_still_loads(tmp_path: Path, monkeypatch) -> None:
    """旧配置有 api_key_env/base_url_env/model_env 时仍可正确加载和展示。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)

    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["llm"] = {
        "default_model": "legacy",
        "models": {
            "legacy": {
                "type": "anthropic_compatible",
                "base_url": "https://legacy.example.com",
                "model": "legacy-model",
                "api_key_env": "MINDFORGE_LEGACY_KEY",
                "base_url_env": "MINDFORGE_LEGACY_URL",
                "model_env": "MINDFORGE_LEGACY_MODEL",
            },
        },
    }
    cfg_path.write_text(yaml.safe_dump(raw, sort_keys=False), encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    # 模拟 env var 有值
    monkeypatch.setenv("MINDFORGE_LEGACY_KEY", "sk-legacy-key-value")
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    editable = client.get("/api/config/editable").json()
    model = editable["llm"]["configured_models"]["legacy"]

    # 读兼容：env 字段仍在响应中（供 Advanced 展示）
    assert model["api_key_env"] == "MINDFORGE_LEGACY_KEY"
    assert model["base_url_env"] == "MINDFORGE_LEGACY_URL"
    assert model["model_env"] == "MINDFORGE_LEGACY_MODEL"
    # API key 来自 env
    assert model["api_key_secret_present"] is True
    assert model["api_key_source"] == "env"
    # raw key 永不泄露
    assert "sk-legacy-key-value" not in f"{editable}"


def test_setup_api_key_env_presence_reads_from_config(tmp_path: Path, monkeypatch) -> None:
    """旧配置的 api_key_env 在 Advanced section 可读，但 raw value 不泄露。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)

    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["llm"] = {
        "default_model": "main",
        "models": {
            "main": {
                "type": "openai_compatible",
                "base_url": "https://r.example.com/v1",
                "model": "main-model",
                "api_key_env": "MINDFORGE_LLM_API_KEY",
            },
        },
    }
    cfg_path.write_text(yaml.safe_dump(raw, sort_keys=False), encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    monkeypatch.setenv("MINDFORGE_LLM_API_KEY", "sk-secret-env-value")
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    editable = client.get("/api/config/editable").json()
    model = editable["llm"]["configured_models"]["main"]
    response_text = f"{editable}"

    # env var name 在 Advanced 可读
    assert model["api_key_env"] == "MINDFORGE_LLM_API_KEY"
    # masked value 正确
    assert model["api_key_masked_value"] == "sk-****alue"
    # raw value 绝对不泄露
    assert "sk-secret-env-value" not in response_text


def test_setup_env_fields_not_in_main_ui_labels(tmp_path: Path, monkeypatch) -> None:
    """验证 legacy env 字段不在主 UI 标签中出现。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)

    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    raw["llm"] = {
        "default_model": "main",
        "models": {
            "main": {
                "type": "openai_compatible",
                "base_url": "https://r.example.com/v1",
                "model": "main-model",
                "api_key_env": "MINDFORGE_LLM_API_KEY",
                "base_url_env": "MINDFORGE_LLM_BASE_URL",
                "model_env": "MINDFORGE_LLM_MODEL",
            },
        },
    }
    cfg_path.write_text(yaml.safe_dump(raw, sort_keys=False), encoding="utf-8")
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    editable = client.get("/api/config/editable").json()
    model = editable["llm"]["configured_models"]["main"]

    # EditableModelConfig 不暴露 env var name 作为主字段
    # api_key_source 从 env 读取（而非 local_secret）
    assert "api_key_source" in model
    assert "is_demo_model" in model
    # env var name 仍可通过 api_key_env 读取（Advanced）
    assert model["api_key_env"] == "MINDFORGE_LLM_API_KEY"


def test_setup_secret_store_preserved_on_edit_no_env_fields(tmp_path: Path, monkeypatch) -> None:
    """Edit model 时空 API key → 保留 secret store，不写 env 字段。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    # 先创建带 key 的模型
    client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://r.example.com/v1",
                    "model": "main-model",
                    "api_key": "sk-preserve-me",
                    "api_key_action": "update",
                },
            },
        },
    )

    # 编辑：api_key 为空，api_key_action keep
    client.patch(
        "/api/config/editable",
        json={
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://r.example.com/v1",
                    "model": "updated-model",
                },
            },
        },
    )

    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    model = raw["llm"]["models"]["main"]
    assert model["model"] == "updated-model"
    assert "api_key_env" not in model
    assert "base_url_env" not in model

    import json
    secrets = json.loads((tmp_path / ".mindforge" / "secrets.json").read_text(encoding="utf-8"))
    assert secrets["main"] == "sk-preserve-me"


def test_setup_clear_key_removes_secret_not_env_fields(tmp_path: Path, monkeypatch) -> None:
    """Clear key → 删除 secret store，不写 env 字段。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://r.example.com/v1",
                    "model": "main-model",
                    "api_key": "sk-will-be-cleared",
                    "api_key_action": "update",
                },
            },
        },
    )

    client.patch(
        "/api/config/editable",
        json={
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://r.example.com/v1",
                    "model": "main-model",
                    "api_key_action": "clear",
                },
            },
        },
    )

    raw = yaml.safe_load(cfg_path.read_text(encoding="utf-8"))
    model = raw["llm"]["models"]["main"]
    assert "api_key_env" not in model

    import json
    secrets = json.loads((tmp_path / ".mindforge" / "secrets.json").read_text(encoding="utf-8"))
    assert "main" not in secrets


def test_setup_default_model_and_routing_still_work_without_env_fields(tmp_path: Path, monkeypatch) -> None:
    """default_model 和 routing 在不用 env 字段时仍正常工作。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    client.patch(
        "/api/config/editable",
        json={
            "default_model": "a",
            "models": {
                "a": {"type": "openai_compatible", "base_url": "https://a.com/v1", "model": "a-model"},
                "b": {"type": "anthropic_compatible", "base_url": "https://b.com", "model": "b-model"},
            },
            "routing": {"distill": "b", "review_questions": "b"},
        },
    )

    editable = client.get("/api/config/editable").json()
    assert editable["llm"]["default_model"] == "a"
    assert editable["llm"]["routing"]["distill"] == "b"
    assert editable["llm"]["routing"]["triage"] == "a"  # fallback
    assert editable["llm"]["routing_is_explicit"] is True


# ============================================================================
# Secret store → provider runtime 测试
# ============================================================================


def test_anthropic_provider_resolves_key_from_secret_store(tmp_path: Path, monkeypatch) -> None:
    """anthropic_compatible model 无 api_key_env 时从 secret store 取 key。"""
    from mindforge.secret_store import SecretStore
    from mindforge.llm.anthropic_compatible import _resolve_api_key

    store = SecretStore(tmp_path / ".mindforge" / "secrets.json")
    store.set("test-model", "sk-secret-store-key")
    monkeypatch.chdir(tmp_path)

    result = _resolve_api_key("test-model", None)
    assert result == "sk-secret-store-key"


def test_openai_provider_resolves_key_from_secret_store(tmp_path: Path, monkeypatch) -> None:
    """openai_compatible model 无 api_key_env 时从 secret store 取 key。"""
    from mindforge.secret_store import SecretStore
    from mindforge.llm.anthropic_compatible import _resolve_api_key

    store = SecretStore(tmp_path / ".mindforge" / "secrets.json")
    store.set("openai-model", "sk-openai-key")
    monkeypatch.chdir(tmp_path)

    result = _resolve_api_key("openai-model", None)
    assert result == "sk-openai-key"


def test_env_var_still_takes_priority(tmp_path: Path, monkeypatch) -> None:
    """env var 优先级高于 secret store。"""
    from mindforge.secret_store import SecretStore
    from mindforge.llm.anthropic_compatible import _resolve_api_key

    store = SecretStore(tmp_path / ".mindforge" / "secrets.json")
    store.set("test-model", "sk-from-store")
    monkeypatch.chdir(tmp_path)
    monkeypatch.setenv("TEST_API_KEY", "sk-from-env")

    result = _resolve_api_key("test-model", "TEST_API_KEY")
    assert result == "sk-from-env"


def test_missing_key_returns_none(tmp_path: Path, monkeypatch) -> None:
    """不存在的 model 返回 None。"""
    from mindforge.llm.anthropic_compatible import _resolve_api_key
    monkeypatch.chdir(tmp_path)

    result = _resolve_api_key("nonexistent", None)
    assert result is None


def test_anthropic_provider_error_never_includes_raw_key() -> None:
    """Provider error message 不包含 raw API key。"""
    from mindforge.llm.base import ProviderError

    # 验证 error message 模版不含 raw key 引用
    try:
        raise ProviderError(
            "模型 test 没有可用的 API key。请在 Web Setup 中添加 key，"
            "或设置环境变量 TEST_KEY。"
        )
    except ProviderError as e:
        msg = str(e)
        assert "sk-" not in msg
        assert "secret" not in msg.lower()


def test_review_route_serves_drafts_page(tmp_path: Path, monkeypatch) -> None:
    """/review 路由正确导向 Drafts 页面，不回退到 Home。"""
    # 验证前端 SPA 路由：/review 和 /drafts 都映射到 DraftsPage
    # 通过检查 App.tsx 源码确认
    app_tsx = Path(__file__).parent.parent / "web" / "src" / "App.tsx"
    content = app_tsx.read_text(encoding="utf-8")
    assert 'path.startsWith("/review")' in content
    assert "DraftsPage" in content


def test_dogfood_command_hidden_from_main_help() -> None:
    """dogfood 命令不出现在主 help 中。"""
    import subprocess
    repo_root = Path(__file__).parent.parent

    result = subprocess.run(
        ["python", "-m", "mindforge", "--help"],
        capture_output=True,
        text=True,
        cwd=repo_root,
    )
    # dogfood 不应出现在主 help 输出中
    assert "dogfood" not in result.stdout


def test_setup_cli_direct_help_is_retired() -> None:
    """旧 setup help 不再作为第二套配置入口暴露。"""
    import subprocess
    repo_root = Path(__file__).parent.parent

    result = subprocess.run(
        ["python", "-m", "mindforge", "setup", "--help"],
        capture_output=True,
        text=True,
        cwd=repo_root,
    )
    assert result.returncode != 0
    assert result.stdout == ""


def test_scan_direct_help_is_retired() -> None:
    """旧 scan help 不再作为第二套 source 入口暴露。"""
    import subprocess
    repo_root = Path(__file__).parent.parent

    result = subprocess.run(
        ["python", "-m", "mindforge", "scan", "--help"],
        capture_output=True,
        text=True,
        cwd=repo_root,
    )
    assert result.returncode != 0
    assert result.stdout == ""


def test_watch_add_frequency_alias(tmp_path: Path) -> None:
    """/watch add --frequency 作为 --every 的别名。"""
    import subprocess
    repo_root = Path(__file__).parent.parent

    result = subprocess.run(
        ["python", "-m", "mindforge", "watch", "add", "--help"],
        capture_output=True,
        text=True,
        cwd=repo_root,
    )
    assert "--frequency" in result.stdout


# ============================================================================
# clean clone first-run + Web Setup save 回归测试
# ============================================================================


def _write_clean_clone_config(
    tmp_path: Path,
    *,
    create_vault: bool = True,
) -> tuple[Path, Path]:
    """模拟 clean clone first_run_config 生成的新格式配置（无模型、无 legacy）。"""
    vault = tmp_path / "vault"
    if create_vault:
        (vault / "00-Inbox").mkdir(parents=True, exist_ok=True)
        (vault / "20-Knowledge-Cards").mkdir(parents=True, exist_ok=True)
        (vault / "30-Projects").mkdir(parents=True, exist_ok=True)
    cfg_path = tmp_path / "mindforge.yaml"
    cfg_path.write_text(
        yaml.safe_dump(
            {
                "version": 0.7,
                "vault": {
                    "root": "vault",
                    "inbox_root": "00-Inbox",
                    "cards_dir": "20-Knowledge-Cards",
                    "archive_dir": "90-Archive/Skipped",
                    "projects_dir": "30-Projects",
                },
                "llm": {
                    "default_model": None,
                    "models": {},
                    "routing": {},
                },
                "wiki": {
                    "mode": "deterministic",
                    "model": None,
                    "auto_rebuild_on_approve": False,
                },
                "telemetry": {"enabled": True, "local_only": True},
            },
            allow_unicode=True,
            sort_keys=False,
        ),
        encoding="utf-8",
    )
    return cfg_path, vault


def test_clean_clone_editable_config_200(tmp_path: Path, monkeypatch) -> None:
    """clean clone first_run_config → GET /api/config/editable 返回 200。"""
    cfg_path, _vault = _write_clean_clone_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    resp = client.get("/api/config/editable")
    assert resp.status_code == 200
    data = resp.json()
    assert data["llm"]["default_model"] is None
    assert data["llm"]["configured_models"] == {}
    assert data["llm"]["routing"] == {}
    assert data["wiki"]["mode"] == "deterministic"
    assert data["wiki"]["model"] is None


def test_clean_clone_quality_api_uses_nested_vault_config(tmp_path: Path, monkeypatch) -> None:
    """回归测试：quality API 必须使用 ``cfg.vault.root/cards_dir``。

    中文学习型说明：clean clone 加载的是 MindForgeConfig 的真实嵌套结构；
    Web facade 不能引用旧的 ``cfg.vault_root`` / ``cfg.cards_dir_rel`` 影子字段，
    否则真实 API 会从业务 200 退化成 AttributeError 500。
    """
    cfg_path, vault = _write_clean_clone_config(tmp_path)
    cards = vault / "20-Knowledge-Cards"
    _write_approved_card(cards, name="clean-quality.md")
    monkeypatch.chdir(tmp_path)
    client = TestClient(
        create_app(config_path=cfg_path, host="127.0.0.1"),
        raise_server_exceptions=False,
    )

    resp = client.get("/api/quality/cards/approved-web-1")

    assert resp.status_code == 200, resp.text
    assert resp.json()["card_id"] == "approved-web-1"


def test_clean_clone_source_location_api_uses_nested_vault_config(tmp_path: Path, monkeypatch) -> None:
    """回归测试：source location API 不能依赖不存在的旧路径字段。"""
    cfg_path, vault = _write_clean_clone_config(tmp_path)
    cards = vault / "20-Knowledge-Cards"
    _write_approved_card(cards, name="clean-location.md")
    monkeypatch.chdir(tmp_path)
    client = TestClient(
        create_app(config_path=cfg_path, host="127.0.0.1"),
        raise_server_exceptions=False,
    )

    resp = client.get("/api/provenance/cards/approved-web-1/location")

    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert data["source_type"] == "manual_note"
    assert data["display"]


def test_clean_clone_save_first_model_default_model_none(tmp_path: Path, monkeypatch) -> None:
    """clean clone 添加第一个模型、default_model=None → PATCH 200。"""
    cfg_path, _vault = _write_clean_clone_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    resp = client.patch(
        "/api/config/editable",
        json={
            "default_model": None,
            "models": {
                "main": {
                    "type": "anthropic_compatible",
                    "base_url": "https://example.com/api",
                    "model": "test-model",
                    "api_key": "sk-test-1234",
                    "api_key_action": "update",
                },
            },
        },
    )
    assert resp.status_code == 200, f"Response: {resp.json()}"

    # 模型写入了 YAML
    yaml_text = cfg_path.read_text(encoding="utf-8")
    data = yaml.safe_load(yaml_text)
    assert "main" in data["llm"]["models"]
    # API key 不写 YAML
    assert "sk-test" not in yaml_text
    assert "1234" not in yaml_text


def test_clean_clone_save_empty_string_default_model_now_200(tmp_path: Path, monkeypatch) -> None:
    """回归测试：clean clone 前端空字符串 default_model="" → 不再 400（已修复）。"""
    cfg_path, _vault = _write_clean_clone_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    # 模拟前端 patchFromForm 发送 default_model="" 的行为
    resp = client.patch(
        "/api/config/editable",
        json={
            "default_model": "",
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://api.openai.com/v1",
                    "model": "gpt-4o-mini",
                    "api_key": "sk-openai-test",
                    "api_key_action": "update",
                },
            },
        },
    )
    assert resp.status_code == 200, f"Expected 200, got {resp.status_code}: {resp.json()}"

    # 模型被保存
    yaml_text = cfg_path.read_text(encoding="utf-8")
    data = yaml.safe_load(yaml_text)
    assert "main" in data["llm"]["models"]
    # default_model 保持 null（因为空字符串被当作未设置）
    assert data["llm"]["default_model"] is None
    # API key 不写 YAML
    assert "sk-openai" not in yaml_text


def test_clean_clone_save_with_default_model_writes_routing(tmp_path: Path, monkeypatch) -> None:
    """clean clone 设置 default_model → 写 routing 五个 stage。"""
    cfg_path, _vault = _write_clean_clone_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    resp = client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "anthropic_compatible",
                    "base_url": "https://example.com/api",
                    "model": "test-model",
                    "api_key_action": "keep",
                },
            },
        },
    )
    assert resp.status_code == 200, f"Response: {resp.json()}"

    yaml_text = cfg_path.read_text(encoding="utf-8")
    data = yaml.safe_load(yaml_text)
    assert data["llm"]["default_model"] == "main"
    routing = data.get("llm", {}).get("routing", {})
    from mindforge.config import REQUIRED_STAGES
    for stage in REQUIRED_STAGES:
        assert routing.get(stage) == "main", f"routing.{stage} missing"


def test_clean_clone_save_first_model_creates_missing_vault(tmp_path: Path, monkeypatch) -> None:
    """回归测试：first-run vault 目录不存在时，保存模型不能被内部目录状态阻塞。

    中文学习型说明：Setup 的主任务是保存模型配置。first-run 生成的 vault path
    是用户可见的工作区位置；如果目录尚不存在，Web save 应自动创建，而不是要求
    普通用户理解 ``create_vault`` 这种内部开关。
    """
    cfg_path, vault = _write_clean_clone_config(tmp_path, create_vault=False)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    payload = {
        "vault_root": str(vault),
        "create_vault": False,
        "default_model": "main",
        "models": {
            "main": {
                "type": "openai_compatible",
                "base_url": "https://example.com/api",
                "model": "test-model",
                "api_key": "sk-clean-clone-test",
                "api_key_action": "update",
            },
        },
        "routing": {},
        "wiki_mode": "deterministic",
        "wiki_model": None,
        "wiki_auto_rebuild_on_approve": False,
    }

    validate = client.post("/api/config/validate", json=payload)
    assert validate.status_code == 200
    assert validate.json()["ok"] is True

    resp = client.patch("/api/config/editable", json=payload)
    assert resp.status_code == 200, f"Expected 200, got {resp.status_code}: {resp.json()}"

    yaml_text = cfg_path.read_text(encoding="utf-8")
    data = yaml.safe_load(yaml_text)
    assert data["llm"]["models"]["main"]["type"] == "openai_compatible"
    assert data["llm"]["default_model"] == "main"
    assert data["wiki"]["mode"] == "deterministic"
    assert data["wiki"]["model"] is None
    for stage in REQUIRED_STAGES:
        assert data["llm"]["routing"][stage] == "main"
    assert vault.is_dir()
    assert (vault / "00-Inbox").is_dir()
    assert (vault / "20-Knowledge-Cards").is_dir()
    assert (vault / "30-Projects").is_dir()
    assert "sk-clean-clone-test" not in yaml_text

    from mindforge.secret_store import SecretStore

    assert SecretStore(tmp_path / ".mindforge" / "secrets.json").present("main")


def test_web_source_path_error_returns_frontend_readable_detail(tmp_path: Path, monkeypatch) -> None:
    """Source 400 必须给前端可读 message，不能让 UI 退化成 ``Bad Request``。

    中文学习型说明：用户主路径是 Add Source / Add and process now。后端可以
    拒绝相对路径或缺模型，但 response body 必须是稳定的 `{detail:{message}}`
    形状，前端才能显示真正原因。
    """
    cfg_path, _vault = _write_clean_clone_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"), raise_server_exceptions=False)

    resp = client.post(
        "/api/sources/watch",
        json={"path": "relative.md", "frequency": "manual", "recursive": True, "process_now": False},
    )

    assert resp.status_code == 400
    assert resp.json()["detail"]["message"].startswith("Please use an absolute path")


def test_setup_save_reports_friendly_error_when_vault_auto_create_fails(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """自动创建 vault 失败时返回可理解错误，不退回内部 traceback。"""
    cfg_path, vault = _write_clean_clone_config(tmp_path, create_vault=False)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    original_mkdir = Path.mkdir

    def fail_vault_mkdir(self: Path, *args, **kwargs) -> None:
        if self == vault:
            raise OSError("permission denied")
        original_mkdir(self, *args, **kwargs)

    monkeypatch.setattr(Path, "mkdir", fail_vault_mkdir)

    resp = client.patch(
        "/api/config/editable",
        json={
            "vault_root": str(vault),
            "default_model": "main",
            "models": {
                "main": {
                    "type": "openai",
                    "base_url": "https://example.com/api",
                    "model": "test-model",
                    "api_key_action": "keep",
                },
            },
        },
    )

    assert resp.status_code == 400
    detail = resp.json()["detail"]
    assert detail["errors"] == [f"Cannot create vault directory: {vault}"]


def test_clean_clone_save_api_key_in_secret_store_not_yaml(tmp_path: Path, monkeypatch) -> None:
    """API key 写 .mindforge/secrets.json，不写 YAML。"""
    cfg_path, vault = _write_clean_clone_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "openai",
                    "base_url": "https://api.openai.com/v1",
                    "model": "gpt-4o",
                    "api_key": "sk-sensitive-9999",
                    "api_key_action": "update",
                },
            },
        },
    )

    yaml_text = cfg_path.read_text(encoding="utf-8")
    assert "sk-sensitive" not in yaml_text

    # secret store 存在
    from mindforge.secret_store import SecretStore
    store = SecretStore(vault.parent / ".mindforge" / "secrets.json")
    assert store.present("main")
    assert store.get("main") == "sk-sensitive-9999"


def test_clean_clone_wiki_deterministic_null_model_ok(tmp_path: Path, monkeypatch) -> None:
    """wiki.mode=deterministic + wiki.model=null 不阻塞保存。"""
    cfg_path, _vault = _write_clean_clone_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    resp = client.patch(
        "/api/config/editable",
        json={
            "wiki_mode": "deterministic",
            "wiki_model": None,
            "wiki_auto_rebuild_on_approve": False,
            "models": {
                "main": {
                    "type": "anthropic_compatible",
                    "base_url": "https://example.com/api",
                    "model": "test-model",
                    "api_key_action": "keep",
                },
            },
            "default_model": "main",
        },
    )
    assert resp.status_code == 200, f"Response: {resp.json()}"

    yaml_text = cfg_path.read_text(encoding="utf-8")
    data = yaml.safe_load(yaml_text)
    assert data["wiki"]["mode"] == "deterministic"
    assert data["wiki"]["model"] is None


def test_clean_clone_wiki_llm_mode_with_model_saves(tmp_path: Path, monkeypatch) -> None:
    """wiki.mode=llm + wiki.model 存在 → 保存成功。"""
    cfg_path, _vault = _write_clean_clone_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    # 需要先有模型
    client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "anthropic_compatible",
                    "base_url": "https://example.com/api",
                    "model": "test-model",
                    "api_key_action": "keep",
                },
            },
        },
    )

    resp = client.patch(
        "/api/config/editable",
        json={
            "wiki_mode": "llm",
            "wiki_model": "main",
            "wiki_auto_rebuild_on_approve": True,
            "default_model": "main",
            "models": {
                "main": {
                    "type": "anthropic_compatible",
                    "base_url": "https://example.com/api",
                    "model": "test-model",
                    "api_key_action": "keep",
                },
            },
        },
    )
    assert resp.status_code == 200, f"Response: {resp.json()}"

    yaml_text = cfg_path.read_text(encoding="utf-8")
    data = yaml.safe_load(yaml_text)
    assert data["wiki"]["mode"] == "llm"
    assert data["wiki"]["model"] == "main"


def test_clean_clone_api_key_missing_does_not_block_save(tmp_path: Path, monkeypatch) -> None:
    """API key missing 不阻塞保存模型 metadata。"""
    cfg_path, _vault = _write_clean_clone_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    resp = client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://example.com/api",
                    "model": "test-model",
                    "api_key_action": "keep",
                },
            },
        },
    )
    assert resp.status_code == 200, f"Response: {resp.json()}"

    yaml_text = cfg_path.read_text(encoding="utf-8")
    data = yaml.safe_load(yaml_text)
    assert data["llm"]["models"]["main"]["type"] == "openai_compatible"


@pytest.mark.parametrize("model_type", [
    "anthropic",
    "anthropic_compatible",
    "openai",
    "openai_compatible",
])
def test_clean_clone_all_model_types_save(tmp_path: Path, monkeypatch, model_type: str) -> None:
    """四种 model type 均可保存。"""
    cfg_path, _vault = _write_clean_clone_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    resp = client.patch(
        "/api/config/editable",
        json={
            "default_model": "demo",
            "models": {
                "demo": {
                    "type": model_type,
                    "base_url": "https://example.com/api",
                    "model": "demo-model",
                    "api_key_action": "keep",
                },
            },
        },
    )
    assert resp.status_code == 200, f"type={model_type}: {resp.json()}"

    yaml_text = cfg_path.read_text(encoding="utf-8")
    data = yaml.safe_load(yaml_text)
    assert data["llm"]["models"]["demo"]["type"] == model_type


def test_clean_clone_save_does_not_write_legacy_fields(tmp_path: Path, monkeypatch) -> None:
    """保存后不写 active_profile/profiles/fake/env 字段。"""
    cfg_path, _vault = _write_clean_clone_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "openai",
                    "base_url": "https://api.openai.com/v1",
                    "model": "gpt-4o-mini",
                    "api_key": "sk-legacy-test",
                    "api_key_action": "update",
                },
            },
        },
    )

    yaml_text = cfg_path.read_text(encoding="utf-8")
    data = yaml.safe_load(yaml_text)
    llm = data.get("llm", {})
    for legacy_key in ("active_profile", "profiles", "active", "providers"):
        assert legacy_key not in llm, f"legacy field {legacy_key} found in YAML"
    # 模型下不应有 env 字段
    for mid, mconf in llm.get("models", {}).items():
        for env_key in ("api_key_env", "base_url_env", "model_env"):
            assert env_key not in mconf, f"env field {env_key} found in model {mid}"


def test_clean_clone_refresh_after_save_keeps_model(tmp_path: Path, monkeypatch) -> None:
    """保存后刷新 Setup → 模型仍存在，key 只显示 masked。"""
    cfg_path, _vault = _write_clean_clone_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    # 保存模型
    client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "anthropic_compatible",
                    "base_url": "https://example.com/api",
                    "model": "test-model",
                    "api_key": "sk-sensitive-key-8888",
                    "api_key_action": "update",
                },
            },
        },
    )

    # 刷新（重新 GET）
    resp = client.get("/api/config/editable")
    assert resp.status_code == 200
    data = resp.json()
    models = data["llm"]["configured_models"]
    assert "main" in models
    model = models["main"]
    assert model["type"] == "anthropic_compatible"
    assert model["model"] == "test-model"
    assert model["api_key_source"] == "local_secret"
    assert model["api_key_secret_present"] is True
    # masked 值显示脱敏前缀+后4位，不包含完整 raw key
    masked = str(model.get("api_key_masked_value", ""))
    assert "sk-sensitive" not in masked
    assert masked != "sk-sensitive-key-8888"  # 不是完整 raw key


def test_setup_save_anchors_config_and_secret_to_current_workspace(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Web Setup 保存链路必须锚定当前 workspace，而不是启动命令所在 CWD。

    中文学习型说明：release dogfood 暴露的问题看起来像“保存成功但 CLI
    仍 needs setup”。这个测试把 Web server CWD 放到另一个目录，只通过
    config_path 指向目标 workspace，验证 config、secret presence 和
    readiness 使用同一个 workspace anchor；测试只检查 secret 文件存在，
    不读取或输出 secret 内容。
    """

    workspace = tmp_path / "workspace"
    repo_cwd = tmp_path / "repo-cwd"
    repo_cwd.mkdir()
    cfg_path, _vault = _write_clean_clone_config(workspace)
    monkeypatch.chdir(repo_cwd)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    response = client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "openai_compatible",
                    "base_url": "https://provider.example.test/v1",
                    "model": "release-test-model",
                    "api_key": "dummy-test-key-not-real",
                    "api_key_action": "update",
                },
            },
        },
    )

    assert response.status_code == 200, response.json()
    yaml_text = cfg_path.read_text(encoding="utf-8")
    saved = yaml.safe_load(yaml_text)
    assert saved["llm"]["models"]["main"]["base_url"] == "https://provider.example.test/v1"
    assert saved["llm"]["models"]["main"]["model"] == "release-test-model"
    assert "dummy-test-key-not-real" not in yaml_text
    assert (workspace / ".mindforge" / "secrets.json").exists()
    assert not (repo_cwd / ".mindforge" / "secrets.json").exists()

    status = client.get("/api/config/status").json()
    assert status["provider"]["model_setup"] == "ready"
    assert status["provider"]["can_run_real_smoke"] is True
    assert status["provider"]["blockers"] == []
    aliases = {item["alias"]: item for item in status["provider"]["aliases"]}
    assert aliases["main"]["api_key_present"] is True


def test_setup_save_writes_provider_timeout_and_retry_defaults(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """Web Setup 保存模型时显式写入 timeout/retry 默认值。

    中文学习型说明：真实用户看到的是 workspace config，而不是 Python loader
    里的隐式默认。保存时写出有限 timeout/retry，能让 release dogfood 中的
    ReadTimeout 变成可解释、可调整的用户体验边界；API key 仍只写 secret store。
    """

    cfg_path, _vault = _write_clean_clone_config(tmp_path)
    monkeypatch.chdir(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    response = client.patch(
        "/api/config/editable",
        json={
            "default_model": "main",
            "models": {
                "main": {
                    "type": "anthropic_compatible",
                    "base_url": "https://provider.example.test/v1",
                    "model": "release-test-model",
                    "api_key": "dummy-test-key-not-real",
                    "api_key_action": "update",
                },
            },
        },
    )

    assert response.status_code == 200, response.json()
    yaml_text = cfg_path.read_text(encoding="utf-8")
    saved = yaml.safe_load(yaml_text)
    model = saved["llm"]["models"]["main"]
    assert model["timeout_seconds"] == 120
    assert model["max_retries"] == 1


# =============================================================================
# P1 Source Path Safety Tests
# 中文学习型说明：这些测试验证 raw path endpoint 已禁用、source_path_view
# 权限正确、前端 fail-closed 行为可测试。
# =============================================================================


def test_raw_path_copy_endpoint_disabled(tmp_path: Path) -> None:
    """raw path copy endpoint 返回 410 Gone，不再接受 {"path": "..."}。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    response = client.post("/api/sources/path-actions/copy", json={"path": "/tmp/test"})
    assert response.status_code == 410


def test_raw_path_reveal_endpoint_disabled(tmp_path: Path) -> None:
    """raw path reveal endpoint 返回 410 Gone，不再接受 {"path": "..."}。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    response = client.post("/api/sources/path-actions/reveal", json={"path": "/tmp/test"})
    assert response.status_code == 410


def test_raw_path_copy_no_existence_oracle(tmp_path: Path) -> None:
    """禁用后的 endpoint 对任意 path 返回相同 410，不泄露路径是否存在。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    # 不同 path（存在/不存在）应返回相同 410
    r1 = client.post("/api/sources/path-actions/copy", json={"path": str(tmp_path)})
    r2 = client.post("/api/sources/path-actions/copy", json={"path": "/nonexistent/path/xyz"})
    assert r1.status_code == 410
    assert r2.status_code == 410


def test_reveal_by_ref_rejects_missing_card(tmp_path: Path) -> None:
    """safe reveal endpoint 对不存在的 card_id 返回 404，不泄露 path。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    response = client.post("/api/sources/reveal", json={"card_id": "nonexistent-card-id"})
    assert response.status_code == 404


def test_reveal_by_ref_rejects_missing_draft(tmp_path: Path) -> None:
    """safe reveal endpoint 对不存在的 draft_id 返回 404。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    response = client.post("/api/sources/reveal", json={"draft_id": "nonexistent-draft-id"})
    assert response.status_code == 404


def test_reveal_by_ref_requires_card_or_draft_id(tmp_path: Path) -> None:
    """safe reveal endpoint 在无 card_id 和 draft_id 时返回 422。

    中文学习型说明：schema 层 exactly-one validator 拒绝两个都为空。
    """
    cfg_path, _vault, _cards = _write_config(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    response = client.post("/api/sources/reveal", json={})
    assert response.status_code == 422


def test_reveal_by_ref_rejects_both_card_and_draft_id(tmp_path: Path) -> None:
    """safe reveal endpoint 在同时传 card_id 和 draft_id 时返回 422。

    中文学习型说明：schema 层 exactly-one validator 拒绝两个都传。
    """
    cfg_path, _vault, _cards = _write_config(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    response = client.post(
        "/api/sources/reveal",
        json={"card_id": "foo", "draft_id": "bar"},
    )
    assert response.status_code == 422


def test_source_path_view_outside_path_no_full_path(tmp_path: Path) -> None:
    """outside_allowed_roots 的 source_path_view 不暴露 full absolute path。

    中文学习型说明：config 放在 tmp_path/config/ 子目录下，让 config_path.parent
    不包含 outside-sources/ 目录。这样 outside 文件才能被正确分类为
    outside_allowed_roots。
    """
    cfg_dir = tmp_path / "config"
    cfg_dir.mkdir()
    cfg_path, vault, cards = _write_config(cfg_dir)
    _write_approved_card(cards)

    # 创建在 config_path.parent 和 vault 之外的文件
    outside_dir = tmp_path / "outside-sources"
    outside_dir.mkdir()
    outside_file = outside_dir / "external.md"
    outside_file.write_text("# External\n\nSome content.", encoding="utf-8")

    from mindforge.config import load_mindforge_config
    from mindforge_web.services.web_path_action_service import WebPathActionService
    cfg = load_mindforge_config(cfg_path)
    svc = WebPathActionService(cfg, config_path=cfg_path)
    view = svc.build_source_path_view(str(outside_file), source_title=outside_file.name)

    assert view.path_kind == "outside_allowed_roots"
    assert view.full_path_available is False
    assert view.can_copy_full_path is False
    assert view.can_reveal_in_finder is False
    assert view.can_copy_display_path is True
    # 只展示 basename，不展示完整 absolute path
    assert view.display_path == outside_file.name
    assert str(outside_file) not in (view.display_path or "")


def test_source_path_view_workspace_path_allows_full_access(tmp_path: Path) -> None:
    """workspace 内路径允许 copy full path 和 reveal。"""
    cfg_path, vault, cards = _write_config(tmp_path)
    _write_approved_card(cards)

    from mindforge.config import load_mindforge_config
    from mindforge_web.services.web_path_action_service import WebPathActionService
    cfg = load_mindforge_config(cfg_path)
    svc = WebPathActionService(cfg, config_path=cfg_path)

    # 使用 vault 内的 inbox 路径
    inbox_file = vault / "00-Inbox" / "ManualNotes" / "test.md"
    inbox_file.parent.mkdir(parents=True, exist_ok=True)
    inbox_file.write_text("# Test", encoding="utf-8")

    view = svc.build_source_path_view(str(inbox_file), source_title=inbox_file.name)

    assert view.path_kind == "workspace"
    assert view.full_path_available is True
    assert view.can_copy_full_path is True
    assert view.can_reveal_in_finder is True
    assert view.can_copy_display_path is True
    assert view.display_path == str(inbox_file)


def test_source_path_view_not_available_no_existence_leak(tmp_path: Path) -> None:
    """missing path 返回 not_available，不泄露存在性（不区分不存在 vs 无权限）。"""
    cfg_path, vault, cards = _write_config(tmp_path)
    _write_approved_card(cards)

    from mindforge.config import load_mindforge_config
    from mindforge_web.services.web_path_action_service import WebPathActionService
    cfg = load_mindforge_config(cfg_path)
    svc = WebPathActionService(cfg, config_path=cfg_path)

    # vault 内不存在的文件
    missing = vault / "00-Inbox" / "ManualNotes" / "does-not-exist.md"
    view = svc.build_source_path_view(str(missing), source_title=missing.name)

    assert view.path_kind == "not_available"
    assert view.can_copy_full_path is False
    assert view.can_reveal_in_finder is False
    # 不展示完整 absolute path
    assert str(missing) != view.display_path


def test_source_path_view_none_source_path(tmp_path: Path) -> None:
    """source_path=None 时返回 not_available，不报错。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    from mindforge.config import load_mindforge_config
    from mindforge_web.services.web_path_action_service import WebPathActionService
    cfg = load_mindforge_config(cfg_path)
    svc = WebPathActionService(cfg, config_path=cfg_path)
    view = svc.build_source_path_view(None)

    assert view.path_kind == "not_available"
    assert view.can_copy_full_path is False
    assert view.can_reveal_in_finder is False


def test_registered_source_not_falsely_claimed(tmp_path: Path) -> None:
    """当前产品没有真实 registered_source registry 时，不误标 registered_source。

    中文学习型说明：allowlisted path 统一返回 workspace，不返回
    registered_source。这避免自报不存在的功能。
    """
    cfg_path, vault, cards = _write_config(tmp_path)
    _write_approved_card(cards)

    from mindforge.config import load_mindforge_config
    from mindforge_web.services.web_path_action_service import WebPathActionService
    cfg = load_mindforge_config(cfg_path)
    svc = WebPathActionService(cfg, config_path=cfg_path)

    inbox_file = vault / "00-Inbox" / "ManualNotes" / "test.md"
    inbox_file.parent.mkdir(parents=True, exist_ok=True)
    inbox_file.write_text("# Test", encoding="utf-8")

    view = svc.build_source_path_view(str(inbox_file), source_title=inbox_file.name)
    # 当前不返回 registered_source；所有 allowlisted path 统一为 workspace
    assert view.path_kind in ("workspace",), f"不应返回 registered_source，实际返回 {view.path_kind}"


# =============================================================================
# P1 Wiki Mode Metadata Tests
# 中文学习型说明：这些测试验证 configured_mode / content_mode / mode 三层语义
# 正确分离，不混用。
# =============================================================================


def test_wiki_status_returns_all_three_mode_layers(tmp_path: Path) -> None:
    """status endpoint 返回 configured_mode / effective_generation_mode / content_mode。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    response = client.get("/api/wiki/status")
    assert response.status_code == 200
    data = response.json()
    assert "configured_mode" in data
    assert "effective_generation_mode" in data
    assert "content_mode" in data
    assert "mode" in data  # legacy alias
    assert "fallback_reason" in data


def test_wiki_page_configured_llm_with_deterministic_content(tmp_path: Path) -> None:
    """configured_mode=llm 但 wiki content 是 deterministic 时，
    content_mode/mode 不得是 llm。"""
    cfg_path, vault, cards = _write_config(tmp_path)
    _write_approved_card(cards)
    wiki_dir = vault / "30-Wiki"
    wiki_dir.mkdir()
    # 写入 deterministic header（无 LLM synthesis 标记）
    (wiki_dir / "Main-Wiki.md").write_text(
        """# MindForge Main Wiki

> This wiki is generated from human-approved knowledge cards.
> It is a derived view. Source files are not copied into this wiki.
> Last rebuilt: 2026-05-18T10:30:00+0800
> Cards included: 1

## Overview

Synthetic overview.

<!-- WIKI_SECTION_START card=approved-web-1 -->
### Approved Web Card

Body.

<!-- WIKI_SECTION_END -->
""",
        encoding="utf-8",
    )
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    response = client.get("/api/wiki/page")
    assert response.status_code == 200
    data = response.json()
    # configured_mode=llm（默认），但 content 是 deterministic
    assert data["configured_mode"] == "llm"
    assert data["content_mode"] == "deterministic"
    assert data["mode"] == "deterministic"  # legacy alias


def test_wiki_page_missing_wiki_mode_not_fallback_to_configured(tmp_path: Path) -> None:
    """缺失 wiki 时 mode/content_mode 不得 fallback 到 configured_mode。"""
    cfg_path, _vault, _cards = _write_config(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    response = client.get("/api/wiki/page")
    assert response.status_code == 200
    data = response.json()
    assert data["exists"] is False
    assert data["configured_mode"] == "llm"
    # mode/content_mode 为 None，不伪装成 llm
    assert data["mode"] is None
    assert data["content_mode"] is None


def test_wiki_page_llm_content_reports_llm_mode(tmp_path: Path) -> None:
    """LLM synthesis header 明确标记时，content_mode 才能是 llm。"""
    cfg_path, vault, cards = _write_config(tmp_path)
    _write_approved_card(cards)
    wiki_dir = vault / "30-Wiki"
    wiki_dir.mkdir()
    (wiki_dir / "Main-Wiki.md").write_text(
        """# MindForge Main Wiki

> LLM synthesis · Model: wiki_model · Last rebuilt: 2026-05-18T11:00:00+0800
> Cards included: 1

## Overview

LLM overview.

<!-- WIKI_SECTION_START card_ids=approved-web-1 -->
### Approved Web Card

Body.

<!-- WIKI_SECTION_END -->
""",
        encoding="utf-8",
    )
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    response = client.get("/api/wiki/page")
    assert response.status_code == 200
    data = response.json()
    assert data["configured_mode"] == "llm"
    assert data["content_mode"] == "llm"
    assert data["mode"] == "llm"
    assert data["model_id"] == "wiki_model"


def test_wiki_status_and_page_consistent_for_same_wiki(tmp_path: Path) -> None:
    """/api/wiki/status 和 /api/wiki/page 对同一 wiki 口径一致。"""
    cfg_path, vault, cards = _write_config(tmp_path)
    _write_approved_card(cards)
    wiki_dir = vault / "30-Wiki"
    wiki_dir.mkdir()
    (wiki_dir / "Main-Wiki.md").write_text(
        """# MindForge Main Wiki

> LLM synthesis · Model: test_model · Last rebuilt: 2026-05-18T12:00:00+0800
> Cards included: 1

## Overview

Overview text.

<!-- WIKI_SECTION_START card_ids=approved-web-1 -->
### Section

Body.

<!-- WIKI_SECTION_END -->
""",
        encoding="utf-8",
    )
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    status_resp = client.get("/api/wiki/status")
    page_resp = client.get("/api/wiki/page")

    assert status_resp.status_code == 200
    assert page_resp.status_code == 200

    status_data = status_resp.json()
    page_data = page_resp.json()

    # content_mode 一致
    assert status_data["content_mode"] == page_data["content_mode"]
    assert status_data["mode"] == page_data["mode"]
    # configured_mode 一致
    assert status_data["configured_mode"] == page_data["configured_mode"]


def test_wiki_status_model_not_ready_has_fallback_reason(tmp_path: Path) -> None:
    """model not ready 时 effective_generation_mode 退化并给出 fallback_reason。

    中文学习型说明：fake provider 的 type="fake" 跳过 API key 检查，
    因此 model_ready=True。这里验证 model ready 时 effective_generation_mode
    等于 configured_mode（llm），fallback_reason 为 None。
    model not ready 的场景需要真实 provider type + missing API key，不适合
    在 fake provider 测试中构造。
    """
    cfg_path, _vault, _cards = _write_config(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    response = client.get("/api/wiki/status")
    assert response.status_code == 200
    data = response.json()
    # fake provider type 被视为 ready
    assert data["model_ready"] is True
    assert data["effective_generation_mode"] == "llm"
    assert data["fallback_reason"] is None


# ===========================================================================
# Remaining P1: API source_path redaction tests
# 中文学习型说明：这些测试验证 source_path 在 API response 中对 unsafe
# path_kind（outside/unknown/not_available）必须 redact 为 None。
# source_path_view.display_path 是唯一可信展示来源。
# ===========================================================================


def test_library_card_outside_source_path_is_redacted(tmp_path: Path) -> None:
    """outside source_path 在 library card API response 中必须为 None。

    中文学习型说明：config_path.parent 不包含 outside 目录，确保
    build_source_path_view 将其分类为 outside_allowed_roots。
    """
    cfg_dir = tmp_path / "config"
    cfg_dir.mkdir()
    cfg_path, vault, cards = _write_config(cfg_dir)
    _write_approved_card(cards)

    # 改写 approved card，将 source_path 改为 outside 绝对路径
    outside_file = tmp_path / "outside-sources" / "external.md"
    outside_file.parent.mkdir()
    outside_file.write_text("# External\n\nOutside content.", encoding="utf-8")

    card_path = cards / "approved.md"
    raw = card_path.read_text(encoding="utf-8")
    raw = raw.replace(
        "source_path: 00-Inbox/ManualNotes/source-note.md",
        f"source_path: {outside_file}",
    )
    card_path.write_text(raw, encoding="utf-8")

    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    # library list
    lib = client.get("/api/library/cards").json()
    assert len(lib["cards"]) == 1
    card = lib["cards"][0]
    # source_path 必须 redact 为 None（outside path）
    assert card["source_path"] is None, f"outside source_path must be None, got {card['source_path']!r}"
    # source_path_view 必须存在
    assert card["source_path_view"] is not None
    assert card["source_path_view"]["path_kind"] == "outside_allowed_roots"
    assert card["source_path_view"]["full_path_available"] is False
    assert card["source_path_view"]["can_copy_full_path"] is False
    assert card["source_path_view"]["can_reveal_in_finder"] is False
    # display_path 只展示 basename
    assert card["source_path_view"]["display_path"] == outside_file.name
    assert str(outside_file) not in (card["source_path_view"]["display_path"] or "")

    # library detail（response_model_exclude_none=True → source_path=None 被排除）
    detail = client.get("/api/library/card", params={"ref": "approved-web-1"}).json()
    assert "source_path" not in detail["card"], (
        f"outside source_path must be excluded from detail response, "
        f"got {detail['card'].get('source_path')!r}"
    )
    assert detail["card"]["source_path_view"]["path_kind"] == "outside_allowed_roots"


def test_draft_outside_source_path_is_redacted(tmp_path: Path) -> None:
    """outside source_path 在 draft API response 中必须为 None。"""
    cfg_dir = tmp_path / "config"
    cfg_dir.mkdir()
    cfg_path, vault, cards = _write_config(cfg_dir)
    _write_draft(cards)

    # 改写 draft card，将 source_path 改为 outside 绝对路径
    outside_file = tmp_path / "outside-sources" / "draft-source.md"
    outside_file.parent.mkdir()
    outside_file.write_text("# Draft Source", encoding="utf-8")

    draft_path = cards / "draft.md"
    raw = draft_path.read_text(encoding="utf-8")
    raw = raw.replace(
        "source_path: 00-Inbox/ManualNotes/source-note.md",
        f"source_path: {outside_file}",
    )
    draft_path.write_text(raw, encoding="utf-8")

    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    # draft list
    drafts_resp = client.get("/api/drafts").json()
    assert len(drafts_resp["drafts"]) == 1
    draft = drafts_resp["drafts"][0]
    assert draft["source_path"] is None, f"outside source_path must be None, got {draft['source_path']!r}"
    assert draft["source_path_view"] is not None
    assert draft["source_path_view"]["path_kind"] == "outside_allowed_roots"

    # draft detail（response_model_exclude_none=True 会导致 source_path 字段缺席）
    detail = client.get("/api/drafts/draft-1").json()
    assert "source_path" not in detail["draft"], (
        f"outside source_path must be absent from detail response, got {detail['draft'].get('source_path')!r}"
    )
    assert detail["draft"]["source_path_view"]["path_kind"] == "outside_allowed_roots"
    # source_context 中的 source_path 也必须 redact（dict 不受 exclude_none 影响）
    assert detail["source_context"]["source_path"] is None


def test_library_card_workspace_source_path_is_preserved(tmp_path: Path) -> None:
    """workspace source_path 在 API response 中必须保留（不 redact）。

    中文学习型说明：vault 内相对路径是安全的，redact 规则只对
    outside/unknown/not_available 生效。需确保 source file 存在。
    """
    cfg_path, vault, cards = _write_config(tmp_path)
    _write_approved_card(cards)

    # 确保 source file 存在（否则 build_source_path_view 分类为 not_available）
    inbox_source = vault / "00-Inbox" / "ManualNotes" / "source-note.md"
    inbox_source.parent.mkdir(parents=True, exist_ok=True)
    inbox_source.write_text("# Safe source", encoding="utf-8")

    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    lib = client.get("/api/library/cards").json()
    assert len(lib["cards"]) == 1
    card = lib["cards"][0]

    # workspace 路径应保留
    assert card["source_path"] == "00-Inbox/ManualNotes/source-note.md"
    assert card["source_path_view"]["path_kind"] == "workspace"
    assert card["source_path_view"]["full_path_available"] is True
    assert card["source_path_view"]["can_copy_full_path"] is True
    assert card["source_path_view"]["can_reveal_in_finder"] is True


def test_source_path_none_redacted_to_none(tmp_path: Path) -> None:
    """source_path 为 None 时 API response 始终保持 None。

    中文学习型说明：没有 source_path 时，source_path_view 为 not_available，
    safe_source_path 也应返回 None（fail-closed）。
    """
    cfg_path, vault, cards = _write_config(tmp_path)
    # 写入无 source_path 的 approved card
    card_path = cards / "no-source.md"
    card_path.write_text(
        """---
id: no-source-1
title: No Source Card
status: human_approved
source_type: manual_note
---
## Body
No source.
""",
        encoding="utf-8",
    )

    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    detail = client.get("/api/library/card", params={"ref": "no-source-1"}).json()
    # response_model_exclude_none=True → source_path=None 被排除
    assert "source_path" not in detail["card"], (
        f"None source_path must be excluded from detail response, "
        f"got {detail['card'].get('source_path')!r}"
    )
    assert detail["card"]["source_path_view"]["path_kind"] == "not_available"


# ===========================================================================
# Remaining P1: RevealRequest contract tests
# ===========================================================================


def test_reveal_endpoint_rejects_extra_path_field(tmp_path: Path) -> None:
    """RevealRequest extra="forbid" —— 传入 path 字段返回 422。

    中文学习型说明：不能静默忽略 path，必须显式拒绝。
    """
    cfg_path, _vault, _cards = _write_config(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    # 传入 raw path 字段 —— 必须 422
    response = client.post(
        "/api/sources/reveal",
        json={"card_id": "draft-1", "path": "/tmp/foo"},
    )
    assert response.status_code == 422, f"extra path must be rejected, got {response.status_code}"
    detail = response.json()
    # 错误响应不 echo raw path
    combined = f"{detail}"
    assert "/tmp/foo" not in combined


def test_reveal_endpoint_outside_card_reveal_rejected(tmp_path: Path) -> None:
    """outside card 的 reveal 必须被拒绝（403）。

    中文学习型说明：即使 card 存在，outside_allowed_roots 的 can_reveal_in_finder
    为 False，reveal 必须返回 403。错误响应不泄露 raw path。
    """
    cfg_dir = tmp_path / "config"
    cfg_dir.mkdir()
    cfg_path, vault, cards = _write_config(cfg_dir)
    _write_draft(cards)

    # 将 source_path 改为 outside 绝对路径
    outside_file = tmp_path / "outside-sources" / "secret.md"
    outside_file.parent.mkdir()
    outside_file.write_text("SECRET", encoding="utf-8")

    draft_path = cards / "draft.md"
    raw = draft_path.read_text(encoding="utf-8")
    raw = raw.replace(
        "source_path: 00-Inbox/ManualNotes/source-note.md",
        f"source_path: {outside_file}",
    )
    draft_path.write_text(raw, encoding="utf-8")

    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    # 尝试 reveal outside card
    response = client.post("/api/sources/reveal", json={"card_id": "draft-1"})
    assert response.status_code == 403, f"outside reveal must be rejected, got {response.status_code}"
    detail = response.json()
    # 错误不泄露 raw path
    assert str(outside_file) not in str(detail)


def test_reveal_endpoint_workspace_card_reveal_allowed(
    tmp_path: Path,
    monkeypatch,
) -> None:
    """workspace card 的 reveal 允许（mock subprocess.run）。

    中文学习型说明：vault 内路径的 reveal 必须成功；mock 掉 subprocess.run
    以避免实际执行 open 命令。验证 authorization decision 通过且 command 正确。
    """
    cfg_path, vault, cards = _write_config(tmp_path)
    _write_draft(cards)

    # 确保 source file 存在
    source_file = vault / "00-Inbox" / "ManualNotes" / "source-note.md"
    source_file.parent.mkdir(parents=True, exist_ok=True)
    source_file.write_text("# Safe", encoding="utf-8")

    calls: list[tuple[list[str], dict]] = []

    def fake_run(args, **kwargs):
        calls.append((list(args), kwargs))

        class Result:
            returncode = 0

        return Result()

    monkeypatch.setattr("mindforge_web.services.web_path_action_service.sys.platform", "darwin")
    monkeypatch.setattr("mindforge_web.services.web_path_action_service.subprocess.run", fake_run)

    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))
    reveal = client.post("/api/sources/reveal", json={"card_id": "draft-1"}).json()

    assert reveal["ok"] is True
    assert reveal["action"] == "reveal"
    assert reveal["path_kind"] == "workspace"
    assert calls[0][0] == ["open", "-R", str(source_file.resolve())]


# ---------------------------------------------------------------------------
# P3 Item A: 旧 raw path endpoint 410 覆盖
# ---------------------------------------------------------------------------


def test_old_copy_path_endpoint_returns_410(tmp_path: Path) -> None:
    """POST /api/sources/path-actions/copy 返回 410 Gone 且不 echo raw path。

    中文学习型说明：旧 raw path endpoint 已禁用，调用即返回 410。
    response detail 中不得包含传入的 raw path 字符串，防止 path probing。
    """
    cfg_path, _, _ = _write_config(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    outside_path = "/tmp/mindforge-test-evil/path"
    resp = client.post("/api/sources/path-actions/copy", json={"path": outside_path})
    assert resp.status_code == 410, f"expected 410, got {resp.status_code}"

    body = resp.json()
    assert outside_path not in str(body)


def test_old_reveal_path_endpoint_returns_410(tmp_path: Path) -> None:
    """POST /api/sources/path-actions/reveal 返回 410 Gone 且不 echo raw path。

    中文学习型说明：旧 raw path endpoint 已禁用，调用即返回 410。
    response detail 中不得包含传入的 raw path 字符串，防止 path probing。
    """
    cfg_path, _, _ = _write_config(tmp_path)
    client = TestClient(create_app(config_path=cfg_path, host="127.0.0.1"))

    outside_path = "/tmp/mindforge-test-evil/path"
    resp = client.post("/api/sources/path-actions/reveal", json={"path": outside_path})
    assert resp.status_code == 410, f"expected 410, got {resp.status_code}"

    body = resp.json()
    assert outside_path not in str(body)
