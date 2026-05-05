"""v0.2.5 — vault index/links + PDF/Docx adapter + CLI polish #2 测试。

为什么 vault index/links 是"安全可见层"而非自动改写：
- 仅写入新文件（``_index.md`` / ``_link_candidates.md``）；
- 不修改任何已有 Knowledge Card 正文；
- 文件首行 marker 用于幂等覆盖判定。

为什么 PDF/Docx 仅最小文本抽取：
- 不做 OCR / 表格 / 复杂版式（详见 ``README.md`` 的 adapter 边界）；
- 依赖通过 ``[project.optional-dependencies]`` 声明，未安装时 lazy 报错。

为什么 doctor 不读 .env：
- 内容一旦读出就有泄漏到日志/输出的风险；
- 只检查文件存在 + .gitignore 是否包含，不读 value。
"""

from __future__ import annotations

import json
import re
from pathlib import Path

import pytest
import yaml
from typer.testing import CliRunner

from mindforge.cli import app
from mindforge.sources.docx import DocxAdapter
from mindforge.sources.pdf import OptionalDependencyError, PdfAdapter
from mindforge.vault import (
    INDEX_MARKER,
    LINKS_MARKER,
    build_index_entries,
    build_link_candidates,
    refresh_indexes,
    render_index_markdown,
    write_link_candidates,
)

runner = CliRunner()


# ---------------------------------------------------------------------------
# Helpers — 复用 v0.2.4 的最小 vault 构造，但加几张真实卡片。
# ---------------------------------------------------------------------------


def _write(p: Path, txt: str) -> Path:
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(txt, encoding="utf-8")
    return p


def _make_vault_with_cards(tmp_path: Path) -> Path:
    """构造 vault + 3 张测试卡片 + 配置；返回 cfg path。"""
    vault = tmp_path / "vault"
    cards = vault / "20-Knowledge-Cards"
    (vault / "00-Inbox" / "ManualNotes").mkdir(parents=True)
    (vault / "30-Projects").mkdir(parents=True)
    (vault / "80-Reviews").mkdir(parents=True)
    cards.mkdir(parents=True)
    (vault / "30-Projects" / "alpha.md").write_text(
        "---\nname: alpha\n---\n\n# alpha\n", encoding="utf-8"
    )

    def card(name: str, **fm) -> None:
        front = "\n".join(f"{k}: {json.dumps(v, ensure_ascii=False)}" for k, v in fm.items())
        (cards / f"{name}.md").write_text(
            f"---\n{front}\n---\n\n# {fm.get('title', name)}\n\nbody\n",
            encoding="utf-8",
        )

    card(
        "c1",
        id="c1",
        title="Agent runtime checkpoint",
        status="human_approved",
        track="agent-runtime",
        projects=["alpha"],
        tags=["agent", "runtime"],
        source_type="webclip_markdown",
    )
    card(
        "c2",
        id="c2",
        title="Agent runtime tool calls",
        status="human_approved",
        track="agent-runtime",
        projects=["alpha"],
        tags=["agent", "tools"],
        source_type="cubox_markdown",
    )
    card(
        "c3",
        id="c3",
        title="Stock analysis basics",
        status="ai_draft",
        track="stock-analysis",
        projects=[],
        tags=["finance"],
        source_type="plain_markdown",
    )

    cfg = {
        "version": 0.1,
        "vault": {
            "root": str(vault),
            "inbox_root": "00-Inbox",
            "cards_dir": "20-Knowledge-Cards",
            "archive_dir": "90-Archive/Skipped",
            "projects_dir": "30-Projects",
        },
        "sources": {
            "enabled": ["plain_markdown"],
            "registry": {
                "plain_markdown": {
                    "adapter": "PlainMarkdownAdapter",
                    "inbox_subdir": "ManualNotes",
                    "file_glob": "*.md",
                    "enabled": True,
                }
            },
        },
        "state": {
            "workdir": str(tmp_path / ".mindforge"),
            "state_file": "state.json",
            "runs_dir": "runs",
            "index_file": "index.jsonl",
        },
        "triage": {"value_score_threshold": 5, "default_track": "unrouted"},
        "llm": {
            "active_profile": "fake",
            "profiles": {
                "fake": {
                    "triage": "f1",
                    "distill": "f1",
                    "link_suggestion": "f1",
                    "review_questions": "f1",
                    "action_extraction": "f1",
                }
            },
            "models": {
                "f1": {
                    "provider": "fake-local",
                    "type": "fake",
                    "base_url": "fake://",
                    "model": "fake-1",
                    "timeout_seconds": 5,
                    "max_retries": 0,
                }
            },
        },
        "prompts": {
            "triage_version": "v1",
            "distill_version": "v1",
            "link_suggestion_version": "v1",
            "review_questions_version": "v1",
            "action_extraction_version": "v1",
        },
        "logging": {"level": "INFO", "file": str(tmp_path / "mf.log")},
    }
    cfg_path = tmp_path / "mindforge.yaml"
    cfg_path.write_text(yaml.safe_dump(cfg, allow_unicode=True), encoding="utf-8")
    return cfg_path


# ---------------------------------------------------------------------------
# vault index
# ---------------------------------------------------------------------------


def test_vault_index_writes_idempotent(tmp_path: Path) -> None:
    cfg_path = _make_vault_with_cards(tmp_path)
    res1 = runner.invoke(app, ["vault", "index", "--config", str(cfg_path)])
    assert res1.exit_code == 0, res1.output
    cards_index = tmp_path / "vault" / "20-Knowledge-Cards" / "_index.md"
    assert cards_index.exists()
    content1 = cards_index.read_text(encoding="utf-8")
    assert content1.startswith(INDEX_MARKER)

    # 二次运行：内容应当与首次完全一致（幂等）
    res2 = runner.invoke(app, ["vault", "index", "--config", str(cfg_path)])
    assert res2.exit_code == 0
    assert cards_index.read_text(encoding="utf-8") == content1


def test_vault_index_does_not_overwrite_user_file(tmp_path: Path) -> None:
    """若 _index.md 已存在但不是 MindForge 维护的，应写到 sibling，不覆盖人手内容。"""
    cfg_path = _make_vault_with_cards(tmp_path)
    cards_dir = tmp_path / "vault" / "20-Knowledge-Cards"
    user_file = cards_dir / "_index.md"
    user_file.write_text("# my hand-written index\n", encoding="utf-8")
    runner.invoke(app, ["vault", "index", "--config", str(cfg_path)])
    assert user_file.read_text(encoding="utf-8") == "# my hand-written index\n"
    sibling = cards_dir / "_index.mindforge.md"
    assert sibling.exists() and sibling.read_text(encoding="utf-8").startswith(INDEX_MARKER)


def test_vault_index_no_secrets(tmp_path: Path) -> None:
    cfg_path = _make_vault_with_cards(tmp_path)
    runner.invoke(app, ["vault", "index", "--config", str(cfg_path)])
    text = (tmp_path / "vault" / "20-Knowledge-Cards" / "_index.md").read_text("utf-8")
    for pat in (r"sk-[A-Za-z0-9]{8,}", r"Bearer ", r"Authorization:"):
        assert not re.search(pat, text)
    # 不应包含 raw_text / prompt / completion 字眼
    for forbidden in ("raw_text", "prompt", "completion", "api_key"):
        assert forbidden not in text


def test_render_index_handles_empty() -> None:
    md = render_index_markdown("Empty", build_index_entries([]))
    assert md.startswith(INDEX_MARKER)
    assert "no cards yet" in md


# ---------------------------------------------------------------------------
# vault links
# ---------------------------------------------------------------------------


def test_vault_links_generates_candidates_safely(tmp_path: Path) -> None:
    cfg_path = _make_vault_with_cards(tmp_path)
    res = runner.invoke(app, ["vault", "links", "--config", str(cfg_path)])
    assert res.exit_code == 0, res.output
    links = tmp_path / "vault" / "20-Knowledge-Cards" / "_link_candidates.md"
    assert links.exists()
    text = links.read_text(encoding="utf-8")
    assert text.startswith(LINKS_MARKER)
    # c1 与 c2 应互为候选（同 track + project + 共享 tag agent）
    assert "Agent runtime tool calls" in text
    assert "Agent runtime checkpoint" in text
    # c3 与前两张几乎不应连
    assert "score=" in text


def test_link_candidates_dedup_and_threshold(tmp_path: Path) -> None:
    from mindforge.cards import iter_cards

    _make_vault_with_cards(tmp_path)
    res = iter_cards(tmp_path / "vault", "20-Knowledge-Cards")
    cands = build_link_candidates(res.cards, top_k=5, min_score=999)
    # 阈值很高：应当几乎没有候选
    assert all(len(cc.candidates) == 0 for cc in cands)


def test_vault_links_dry_run_does_not_write(tmp_path: Path) -> None:
    cfg_path = _make_vault_with_cards(tmp_path)
    runner.invoke(app, ["vault", "links", "--config", str(cfg_path), "--dry-run"])
    assert not (tmp_path / "vault" / "20-Knowledge-Cards" / "_link_candidates.md").exists()


def test_vault_links_module_helpers(tmp_path: Path) -> None:
    """直接测 build/write helpers，确保模块外可独立调用。"""
    from mindforge.cards import iter_cards

    _make_vault_with_cards(tmp_path)
    res = iter_cards(tmp_path / "vault", "20-Knowledge-Cards")
    cands = build_link_candidates(res.cards)
    assert any(cc.candidates for cc in cands)
    p, content = write_link_candidates(
        tmp_path / "vault" / "20-Knowledge-Cards", cands, dry_run=True
    )
    assert content.startswith(LINKS_MARKER)
    assert not p.exists()  # dry_run 不写


# ---------------------------------------------------------------------------
# refresh_indexes module-level
# ---------------------------------------------------------------------------


def test_refresh_indexes_writes_three_files(tmp_path: Path) -> None:
    cfg_path = _make_vault_with_cards(tmp_path)
    # CLI 已经测过；这里测模块函数返回值
    from mindforge.config import load_mindforge_config

    cfg = load_mindforge_config(cfg_path)
    res = refresh_indexes(
        cfg.vault.root, cfg.vault.cards_dir, cfg.vault.projects_dir, "80-Reviews"
    )
    paths = {p.name for p in res.written}
    assert "_index.md" in paths
    assert len(res.written) == 3  # cards + projects + reviews


# ---------------------------------------------------------------------------
# PDF/Docx adapter — lazy import error path
# ---------------------------------------------------------------------------


def test_pdf_adapter_raises_optional_dep_when_missing(tmp_path: Path, monkeypatch) -> None:
    """模拟 pypdf 不可 import；load() 必须给出可操作错误。"""
    import sys

    monkeypatch.setitem(sys.modules, "pypdf", None)
    f = _write(tmp_path / "x.pdf", "fake bytes")
    a = PdfAdapter()
    assert a.can_handle(str(f))
    with pytest.raises(OptionalDependencyError, match=r"mindforge\[pdf\]"):
        a.load(str(f))


def test_docx_adapter_raises_optional_dep_when_missing(
    tmp_path: Path, monkeypatch
) -> None:
    import sys

    monkeypatch.setitem(sys.modules, "docx", None)
    f = _write(tmp_path / "x.docx", "fake")
    a = DocxAdapter()
    assert a.can_handle(str(f))
    with pytest.raises(OptionalDependencyError, match=r"mindforge\[docx\]"):
        a.load(str(f))


def test_pdf_adapter_real_extraction_if_available(tmp_path: Path) -> None:
    """若 pypdf 可用，做最小往返；否则跳过。"""
    pytest.importorskip("pypdf")
    import pypdf

    f = tmp_path / "tiny.pdf"
    # 用 pypdf 自己生成一个最小 PDF（一页空白 + 标题 metadata）
    writer = pypdf.PdfWriter()
    writer.add_blank_page(width=200, height=200)
    writer.add_metadata({"/Title": "TinyDoc"})
    with open(f, "wb") as fh:
        writer.write(fh)

    a = PdfAdapter()
    # 空白页没有文本层 → 应抛 PdfNoTextError（v0.2.5 不做 OCR）
    from mindforge.sources.pdf import PdfNoTextError

    with pytest.raises(PdfNoTextError):
        a.load(str(f))


def test_docx_adapter_real_extraction_if_available(tmp_path: Path) -> None:
    pytest.importorskip("docx")
    import docx

    f = tmp_path / "tiny.docx"
    d = docx.Document()
    d.add_paragraph("Hello MindForge.")
    d.add_paragraph("Another paragraph.")
    d.save(str(f))

    a = DocxAdapter()
    doc = a.load(str(f))
    assert doc.source_type == "docx"
    assert "Hello MindForge." in doc.raw_text
    assert doc.metadata["paragraph_count"] == 2
    assert doc.content_hash.startswith("sha256:")


# ---------------------------------------------------------------------------
# CLI polish #2 — --vault override / doctor
# ---------------------------------------------------------------------------


def test_vault_override_changes_resolved_root(tmp_path: Path) -> None:
    cfg_path = _make_vault_with_cards(tmp_path)
    other = tmp_path / "other_vault"
    (other / "20-Knowledge-Cards").mkdir(parents=True)
    res = runner.invoke(
        app, ["--vault", str(other), "version", "--config", str(cfg_path)]
    )
    assert res.exit_code == 0
    # Rich 会按终端宽度换行；剥掉换行后再比对子串
    flat = "".join(res.output.split())
    assert str(other).replace("/", "") in flat.replace("/", "") or str(other) in flat


def test_doctor_runs_without_secrets(tmp_path: Path) -> None:
    cfg_path = _make_vault_with_cards(tmp_path)
    # 写一个含敏感样子内容的 .env，再创建 .gitignore 包含它
    (tmp_path / ".env").write_text("MINDFORGE_LLM_API_KEY=sk-shouldnotleak\n", "utf-8")
    (tmp_path / ".gitignore").write_text(".env\n", "utf-8")
    import os

    cwd = os.getcwd()
    try:
        os.chdir(tmp_path)
        res = runner.invoke(app, ["doctor", "--config", str(cfg_path)])
    finally:
        os.chdir(cwd)
    assert res.exit_code == 0, res.output
    # 严禁泄漏 .env value
    assert "sk-shouldnotleak" not in res.output
    assert "gitignored" in res.output
    assert "MindForge doctor" in res.output


def test_doctor_warns_when_env_not_in_gitignore(tmp_path: Path) -> None:
    cfg_path = _make_vault_with_cards(tmp_path)
    (tmp_path / ".env").write_text("X=1\n", "utf-8")
    # 不创建 .gitignore（或不含 .env）
    import os

    cwd = os.getcwd()
    try:
        os.chdir(tmp_path)
        res = runner.invoke(app, ["doctor", "--config", str(cfg_path)])
    finally:
        os.chdir(cwd)
    assert res.exit_code == 0
    assert "not in .gitignore" in res.output


def test_doctor_handles_missing_config(tmp_path: Path) -> None:
    res = runner.invoke(app, ["doctor", "--config", str(tmp_path / "nope.yaml")])
    assert res.exit_code == 0
    assert "MISSING" in res.output
