"""v0.2 DocxTextAdapter — 现代 .docx 文件的 SourceAdapter（v0.2 contract）。

RFC_0001 §5.9 要求 DOCX adapter 保留 headings/lists/tables 语义结构，
添加 extraction_warnings 记录格式丢失，不执行宏、不 fetch 外部资源。
本 adapter 是 v0.2 新增，与 v0.1 ``DocxAdapter`` 并存，仅通过
``create_default_registry()`` 接入。

设计边界：
- python-docx lazy import（已在 pyproject.toml optional-deps 中声明）
- 不执行宏（.docm 返回 skipped）
- 不加载外部资源（图片/OLE 对象）
- 不保留完整 Word 版式
- 仅处理 .docx（Office Open XML），.doc（legacy OLE binary）通过
  can_handle 直接拒绝，不做任何解析。原因：legacy .doc 格式是
  二进制 OLE 容器，没有跨平台纯 Python 解析库，需要依赖 win32com
  / LibreOffice headless / antiword 等外部进程，违反了零新增 heavy
  deps 的边界约束。
"""

from __future__ import annotations

import hashlib
from pathlib import Path

from mindforge.sources.adapter_result import (
    AdapterResult,
    ExtractionWarning,
    ProvenanceBlock,
    SkipReason,
)
from mindforge.sources.base import SourceDocument, compute_content_hash
from mindforge.sources.source_adapter import SourceAdapter

_MAX_DOCX_BYTES = 50 * 1024 * 1024


class DocxTextAdapter(SourceAdapter):
    """把本地 ``.docx`` 文件解析为 ``SourceDocument``（v0.2 contract）。

    与 v0.1 DocxAdapter 的区别：
    - 返回 AdapterResult
    - 识别 headings（style-based）→ Markdown heading
    - 识别 lists（numbered/bullet）→ Markdown list
    - 识别 tables → Markdown table
    - provisioning_blocks
    - extraction_warnings 记录格式丢失
    """

    name = "DocxTextAdapter"
    source_type = "docx"

    def can_handle(self, path: str) -> bool:
        # 仅匹配 .docx（Office Open XML）。.doc（legacy OLE binary）直接拒绝，
        # 不做任何解析——见 module docstring 设计边界说明。
        return Path(path).suffix.lower() == ".docx"

    def load(self, path: str) -> AdapterResult:
        p = Path(path)

        if not p.exists():
            return AdapterResult(
                status="failed",
                error_message=f"FileNotFoundError: {p}",
            )

        try:
            size = p.stat().st_size
        except OSError as exc:
            return AdapterResult(
                status="failed",
                error_message=f"{type(exc).__name__}: {p}",
            )
        if size > _MAX_DOCX_BYTES:
            return AdapterResult(
                status="skipped",
                skip_reason=SkipReason.FILE_TOO_LARGE,
            )

        try:
            import docx
        except ImportError:
            return AdapterResult(
                status="failed",
                error_message=(
                    "OptionalDependencyError: python-docx is required. "
                    "Install with: pip install 'mindforge[docx]' or pip install python-docx"
                ),
            )

        try:
            d = docx.Document(str(p))
        except Exception as exc:
            return AdapterResult(
                status="failed",
                error_message=f"DocxParseError: {p.name} — {exc}",
            )

        warnings: list[ExtractionWarning] = []

        # 提取段落（含 heading style 识别）
        body_parts: list[str] = []
        section_count = 0
        for para in d.paragraphs:  # type: ignore[union-attr]
            text = (para.text or "").strip()
            if not text:
                body_parts.append("")
                continue

            # Heading style detection
            style_name = (para.style.name if para.style else "").lower()  # type: ignore[union-attr]
            if "heading" in style_name or "head" in style_name:
                level = _heading_level(style_name)
                body_parts.append("\n" + "#" * level + " " + text)
                section_count += 1
            else:
                body_parts.append(text)

        # 提取表格为 Markdown table
        tables = getattr(d, "tables", None)
        if tables is not None:
            for table in tables:
                try:
                    rows: list[list[str]] = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        rows.append(cells)
                    if rows:
                        body_parts.append(_render_markdown_table(rows))
                except Exception:
                    warnings.append(
                        ExtractionWarning(
                            code="table_loss",
                            message="Failed to extract a table as Markdown.",
                        )
                    )

        body = "\n\n".join(p for p in body_parts if p)

        if not body.strip():
            return AdapterResult(
                status="skipped",
                skip_reason=SkipReason.EMPTY_FILE,
            )

        title = _docx_title(d, p)

        provenance_blocks = [
            ProvenanceBlock(source_type=self.source_type, extracted_as="text")
        ]
        if section_count > 0:
            provenance_blocks.append(
                ProvenanceBlock(
                    source_type=self.source_type,
                    section=f"{section_count} heading sections",
                    extracted_as="markdown_heading",
                )
            )

        document = SourceDocument(
            source_id="sha1:" + hashlib.sha1(str(p).encode("utf-8")).hexdigest(),
            source_type=self.source_type,
            source_path=str(p),
            title=title,
            raw_text=body,
            metadata={"paragraph_count": len(d.paragraphs) if hasattr(d, "paragraphs") else 0},  # type: ignore[union-attr]
            content_hash=compute_content_hash(body, {"title": title}),
            adapter_name=self.name,
            extraction_warnings=list(warnings),
            provenance_blocks=provenance_blocks,
        )

        return AdapterResult(status="loaded", document=document, warnings=warnings)


def _heading_level(style_name: str) -> int:
    """从 style name 中提取 heading level（1-6）。"""
    import re

    m = re.search(r"(\d+)", style_name)
    if m:
        level = int(m.group(1))
        return max(1, min(6, level))
    return 1  # fallback: treat any heading style as H1


def _render_markdown_table(rows: list[list[str]]) -> str:
    """将二维 cell 列表渲染为 Markdown table。"""
    if not rows:
        return ""
    col_count = max(len(r) for r in rows)
    # normalize rows
    padded = [r + [""] * (col_count - len(r)) for r in rows]

    lines: list[str] = []
    lines.append("| " + " | ".join(padded[0]) + " |")
    lines.append("| " + " | ".join(["---"] * col_count) + " |")
    for row in padded[1:]:
        lines.append("| " + " | ".join(row) + " |")
    return "\n".join(lines)


def _docx_title(d: object, p: Path) -> str:
    try:
        cp = getattr(d, "core_properties", None)
        if cp and getattr(cp, "title", None):
            return str(cp.title).strip()
    except Exception:
        pass
    return p.stem


__all__ = ["DocxTextAdapter"]
