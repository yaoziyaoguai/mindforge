"""DocxAdapter — Word .docx 的最小段落文本抽取（M5.1 v0.2.5）。

边界：
- 仅取 ``Document.paragraphs`` 的纯文本；不解析样式、表格、批注、修订、
  embedded objects、文本框、SmartArt；
- ``lazy import``：未装 ``python-docx`` 时只在 ``load()`` 报错；
- 不修改原始文件。
"""

from __future__ import annotations

import hashlib
from pathlib import Path

from .base import SourceAdapter, SourceDocument, compute_content_hash
from .pdf import OptionalDependencyError


class DocxAdapter(SourceAdapter):
    name = "DocxAdapter"
    source_type = "docx"  # type: ignore[assignment]

    def can_handle(self, path: str) -> bool:
        return path.lower().endswith(".docx")

    def load(self, path: str) -> SourceDocument:
        try:
            import docx  # type: ignore[import-not-found]
        except ImportError as e:
            raise OptionalDependencyError(
                "DocxAdapter 需要 'python-docx'（可选依赖）。请运行：\n"
                "    pip install 'mindforge[docx]'\n"
                "或:\n"
                "    pip install python-docx\n"
                "详见 docs/M5_1_PDF_DOCX_ADAPTER_PROTOCOL.md。"
            ) from e

        p = Path(path)
        if not p.exists():
            raise FileNotFoundError(f"DocxAdapter: 文件不存在 {p}")

        try:
            d = docx.Document(str(p))
        except Exception as e:
            raise OptionalDependencyError(
                f"DocxAdapter: python-docx 解析失败（{type(e).__name__}）：{e}"
            ) from e

        paras = [para.text for para in d.paragraphs if (para.text or "").strip()]
        body = "\n\n".join(paras)

        # 取首段非空作为 title 兜底（如果 core_properties 没有）
        title = _pick_title(d, p)
        author = _pick_author(d)

        source_id = "sha1:" + hashlib.sha1(path.encode("utf-8")).hexdigest()
        key_meta = {"title": title, "paragraph_count": len(paras)}
        content_hash = compute_content_hash(body, key_meta)

        return SourceDocument(
            source_id=source_id,
            source_type=self.source_type,
            source_path=path,
            title=title,
            author=author,
            source_url=None,
            created_at=None,
            captured_at=None,
            tags=[],
            highlights=[],
            raw_text=body,
            metadata={"paragraph_count": len(paras)},
            content_hash=content_hash,
        )


def _pick_title(d: object, p: Path) -> str:
    try:
        cp = getattr(d, "core_properties", None)
        if cp and getattr(cp, "title", None):
            return str(cp.title).strip()
    except Exception:  # noqa: BLE001
        pass
    return p.stem


def _pick_author(d: object) -> str | None:
    try:
        cp = getattr(d, "core_properties", None)
        if cp and getattr(cp, "author", None):
            return str(cp.author).strip()
    except Exception:  # noqa: BLE001
        pass
    return None


__all__ = ["DocxAdapter"]
